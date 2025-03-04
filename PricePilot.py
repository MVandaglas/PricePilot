import streamlit as st
st.set_page_config(page_icon="🎯",layout="wide")
from streamlit_option_menu import option_menu
import os
import pandas as pd
from PIL import Image
import re
import json
from datetime import datetime, timedelta, date
from st_aggrid import AgGrid, GridOptionsBuilder, JsCode, ColumnsAutoSizeMode, GridUpdateMode, DataReturnMode
import openai
from SAPprijs import sap_prices
from Synonyms import synonym_dict
from Articles import article_table
import difflib
from rapidfuzz import process, fuzz
from io import BytesIO
from PyPDF2 import PdfReader
import extract_msg
import pdfplumber
from functools import partial
from database_setup import create_connection, setup_database
import sqlite3
from http.cookies import SimpleCookie
from simple_salesforce import Salesforce, SalesforceLogin
import time
from docx import Document
import xlsxwriter
import getpass
import requests
from requests.auth import HTTPBasicAuth
from requests_ntlm import HttpNtlmAuth 
from office365.runtime.auth.client_credential import ClientCredential
from office365.sharepoint.client_context import ClientContext
from msal import ConfidentialClientApplication
import jwt
import numpy as np
import tempfile
import pyodbc


# 🔑 Configuratie
CLIENT_ID = st.secrets.get("SP_CLIENTID")
CLIENT_SECRET = st.secrets.get("SP_CLIENTSECRET")
SP_SITE = st.secrets.get("SP_SITE")
TENANT_ID = st.secrets.get("TENANT_ID")
CSV_PATH = st.secrets.get("SP_CSV_SYN")  # Pad naar TestSynoniem.csv in SharePoint
SP_USERNAME = st.secrets.get("SP_USERNAME")
SP_PASSWORD = st.secrets.get("SP_PASSWORD")

# **Verbinding met Azure SQL Server**
def create_connection():
    server = "vdgbullsaiserver.database.windows.net,1433"
    database = "vdgbullsaidb"
    username = SP_USERNAME
    password = SP_PASSWORD

    conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={server};DATABASE={database};Authentication=ActiveDirectoryPassword;UID={username};PWD={password}"
    try:
        conn = pyodbc.connect(conn_str)
        return conn
    except Exception as e:
        st.error(f"Database fout: {e}")
        return None


# Importeer prijsscherpte
if "prijsscherpte_matrix" not in st.session_state:
    # Initialiseer de matrix met standaardwaarden
    st.session_state.prijsscherpte_matrix = pd.DataFrame({
        "Offertebedrag": [0, 5000, 10000, 25000, 50000],  # X-as
        "A": [60, 70, 80, 90, 100],  # Y-as kolommen
        "B": [40, 50, 60, 70, 80],
        "C": [30, 40, 50, 65, 75],
        "D": [10, 25, 45, 60, 65],
    })

st.sidebar.write(f"Laatste update: {time.ctime()}")

# Functie om klantgegevens op te halen uit Salesforce zonder caching
def fetch_salesforce_accounts_direct(sf_connection):
    try:
        # Query voor Salesforce-accounts
        accounts_query = sf_connection.query("""
            SELECT Id, Name, ERP_Number__c
            FROM Account
            WHERE ERP_Number__c != NULL AND Is_Active__c = TRUE
            ORDER BY Name ASC
            LIMIT 6000
        """)
        return accounts_query["records"]
    except Exception as e:
        st.error(f"Fout bij ophalen van Salesforce-accounts: {e}")
        return []

# Salesforce Login Configuratie
SF_USERNAME =  os.getenv("SALESFORCE_USERNAME")
SF_PASSWORD = os.getenv("SALESFORCE_PASSWORD") + os.environ.get("SF_SECURITY_TOKEN")
SF_SECURITY_TOKEN =  os.getenv("SF_SECURITY_TOKEN")
SF_DOMAIN = "test"  # Gebruik 'test' voor Sandbox

if "force_rerun" in st.session_state and st.session_state.force_rerun:
    st.session_state.force_rerun = False  # Zet de trigger uit om oneindige loops te voorkomen
    st.rerun()  # UI herladen zonder dat state verloren gaat


# Verbind met Salesforce
try:
    session_id, instance = SalesforceLogin(
        username=SF_USERNAME,
        password=SF_PASSWORD,
        domain=SF_DOMAIN
    ) #keys in manage app secrets
    sf = Salesforce(instance=instance, session_id=session_id)

except Exception as e:
    st.error(f"Fout bij het verbinden met Salesforce: {e}")

# Haal accounts op als de verbinding geslaagd is
if sf:
    accounts = fetch_salesforce_accounts_direct(sf)
else:
    accounts = []

# Verwerk de accounts als er gegevens beschikbaar zijn
if accounts:
    accounts_df = pd.DataFrame(accounts).drop(columns="attributes", errors="ignore")
    accounts_df.rename(columns={"Name": "Klantnaam", "ERP_Number__c": "Klantnummer"}, inplace=True)
    accounts_df["Klantinfo"] = accounts_df["Klantnummer"] + " - " + accounts_df["Klantnaam"]
else:
    accounts_df = pd.DataFrame(columns=["Klantnaam", "Klantnummer", "Klantinfo"])



# OpenAI API-sleutel instellen
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    st.error("OpenAI API-sleutel ontbreekt. Stel de OPENAI_API_KEY omgevingsvariabele in de Streamlit Cloud-instellingen in.")
else:
    openai.api_key = api_key  # Initialize OpenAI ChatCompletion client


# Zorg ervoor dat de database bij opstarten correct is
setup_database()

# Hard gecodeerde klantgegevens
customer_data = {
    "111111": {"revenue": "50.000 euro", "size": "D"},
    "222222": {"revenue": "140.000 euro", "size": "B"},
    "333333": {"revenue": "600.000 euro", "size": "A"},
    "100007": {"revenue": "141.000 euro", "size": "B"},
}

# Initialiseer offerte DataFrame en klantnummer in sessiestatus
if "offer_df" not in st.session_state:
    st.session_state.offer_df = pd.DataFrame(columns=["Rijnummer", "Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", "Aantal", "RSP", "SAP Prijs", "Handmatige Prijs", "Min_prijs", "M2 p/s", "M2 totaal", "Max_prijs", "Verkoopprijs", "Prijs_backend", "Source"])
if "customer_number" not in st.session_state:
    st.session_state.customer_number = ""
if "loaded_offer_df" not in st.session_state:
    st.session_state.loaded_offer_df = pd.DataFrame(columns=["Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", "Aantal", "RSP", "M2 p/s", "M2 totaal", "Verkoopprijs"])
if "saved_offers" not in st.session_state:
    st.session_state.saved_offers = pd.DataFrame(columns=["Offertenummer", "Klantnummer", "Eindbedrag", "Datum"])
if "selected_rows" not in st.session_state:
    st.session_state.selected_rows = []


# Converteer article_table naar DataFrame
article_table = pd.DataFrame(article_table)

# Streamlit UI-instellingen
# Maak de tabs aan
tab1, tab2, tab3, tab4 = st.tabs(["🎯 Offerte Genereren", "💾 Opgeslagen Offertes", "✨ Beoordeel AI", "⚙️ Beheer"])

with tab4:
    st.subheader("Beheer")

    # **Wachtwoordbeveiliging**
    wachtwoord = st.text_input("Voer het wachtwoord in om toegang te krijgen:", type="password")
    if wachtwoord == "Comex25":
        st.success("Toegang verleend tot de beheertab.")

        
        with st.expander("🔍 Bekijk en beheer actieve synoniemen", expanded=False):       
            # **Maak verbinding met de database**
            conn = create_connection()
            if conn:
                cursor = conn.cursor()
            
                try:
                    # **Controleer of de tabel 'synoniemen' bestaat**
                    cursor.execute("""
                    SELECT TABLE_NAME FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'synoniemen';
                    """)
                    tabel_bestaat = cursor.fetchone()
            
                    if tabel_bestaat:
                        # **Haal de geaccordeerde synoniemen op**
                        cursor.execute("SELECT Artikelnummer, Synoniem FROM synoniemen")
                        synoniemen_data = cursor.fetchall()
                        
                        # **Haal de kolomnamen op**
                        kolomnamen = [desc[0] for desc in cursor.description]
     
                        
                        # **Controleer of er None-waarden zijn**
                        for rij in synoniemen_data:
                            if None in rij:
                              pass
                        
                        # **Converteer tuples naar lijsten**
                        synoniemen_data_lijst = [list(rij) for rij in synoniemen_data]

                        
                        # **Maak DataFrame aan**
                        synoniemen_df = pd.DataFrame(synoniemen_data_lijst, columns=kolomnamen)
                        
            
                        if not synoniemen_df.empty:
                            # **Configureer AgGrid voor Synoniemen**
                            gb = GridOptionsBuilder.from_dataframe(synoniemen_df)
                            gb.configure_selection(selection_mode="multiple", use_checkbox=True)
                            gb.configure_default_column(editable=False)
                            grid_options = gb.build()
            
                            response = AgGrid(
                                synoniemen_df,
                                gridOptions=grid_options,
                                update_mode=GridUpdateMode.SELECTION_CHANGED,
                                fit_columns_on_grid_load=True,
                                theme="material"
                            )
            
                            # **Geselecteerde rijen ophalen**
                            geselecteerde_rijen = response["selected_rows"]
            
                            if st.button("Verwijder geselecteerde synoniemen"):
                                if len(geselecteerde_rijen) > 0:
                                    try:
                                        for rij in geselecteerde_rijen:
                                            # Controleer of rij een dictionary is of een tuple/lijst
                                            if isinstance(rij, dict):
                                                synoniem = rij.get("Synoniem")
                                                artikelnummer = rij.get("Artikelnummer")
                                            elif isinstance(rij, (tuple, list)) and len(rij) == 2:
                                                artikelnummer, synoniem = rij  # Pak waarden uit tuple/lijst
                                            else:
                                                st.warning(f"Ongeldig formaat van rij: {rij}")
                                                continue
                            
                                            if synoniem and artikelnummer:
                                                cursor.execute("""
                                                DELETE FROM synoniemen WHERE Artikelnummer = ? AND Synoniem = ?;
                                                """, (artikelnummer, synoniem))
                            
                                        conn.commit()
                                        st.success("Geselecteerde synoniemen zijn verwijderd uit 'synoniemen'.")
                                    except Exception as e:
                                        st.error(f"Fout bij verwijderen van synoniemen: {e}")
                                else:
                                    st.warning("Selecteer minimaal één rij om te verwijderen.")

            
                except Exception as e:
                    st.error(f"Fout bij ophalen van synoniemen: {e}")
            
                finally:
                    conn.close()
    

    elif wachtwoord:
        st.error("❌ Onjuist wachtwoord. Toegang geweigerd.")
        

# Tab 1: Offerte Genereren
with tab1:
    st.subheader("Offerte Genereren")
    
    if st.session_state.offer_df is not None and not st.session_state.offer_df.empty:
        st.title("Offerteoverzicht")




if st.session_state.offer_df is None or st.session_state.offer_df.empty:
    st.session_state.offer_df = pd.DataFrame(columns=["Rijnummer", "Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", "Aantal", "M2 p/s", "M2 totaal", "RSP", "SAP Prijs", "Handmatige Prijs", "Min_prijs", "Max_prijs", "Verkoopprijs", "Prijs_backend", "Source"])


# Omzetting naar numerieke waarden en lege waarden vervangen door 0
st.session_state.offer_df["M2 totaal"] = pd.to_numeric(st.session_state.offer_df["M2 totaal"], errors='coerce').fillna(0)
st.session_state.offer_df["RSP"] = pd.to_numeric(st.session_state.offer_df["RSP"], errors='coerce').fillna(0)
st.session_state.offer_df["Verkoopprijs"] = pd.to_numeric(st.session_state.offer_df["Verkoopprijs"], errors='coerce')

# Offerte Genereren tab
with tab1:
    
    # Voeg een dropdown toe voor prijsbepaling met een breedte-instelling
    col1, _ = st.columns([1, 7])  # Maak kolommen om breedte te beperken
    with col1:
        prijsbepaling_optie = st.selectbox("Prijsbepaling", ["PricePilot logica", "SAP prijs", "RSP"], key="prijsbepaling", help="Selecteer een methode voor prijsbepaling.")

# Offerte Genereren tab
with tab1:
    def bereken_prijs_backend(df):
        if df is None or not isinstance(df, pd.DataFrame):
            st.warning("De DataFrame is leeg of ongeldig. Prijs_backend kan niet worden berekend.")
            return pd.DataFrame()  # Retourneer een lege DataFrame als fallback

        try:
            # Controleer of de DataFrame geldig is
            if not isinstance(df, pd.DataFrame):
                raise ValueError("De input is geen geldige DataFrame.")

            # Zorg ervoor dat kolommen numeriek zijn of bestaan
            for col in ["SAP Prijs", "RSP", "Handmatige Prijs", "Prijskwaliteit"]:
                if col not in df.columns:
                    df[col] = 0  # Voeg de kolom toe als deze niet bestaat

            df["SAP Prijs"] = pd.to_numeric(df["SAP Prijs"], errors="coerce").fillna(0)
            df["RSP"] = pd.to_numeric(df["RSP"], errors="coerce").fillna(0)
            df["Handmatige Prijs"] = pd.to_numeric(df["Handmatige Prijs"], errors="coerce").fillna(0)
            df["Prijskwaliteit"] = pd.to_numeric(df["Prijskwaliteit"], errors="coerce").fillna(100)

            # Functie om Prijs_backend te bepalen op basis van logica
            def bepaal_prijs_backend(row):
                # Controleer of Handmatige Prijs is ingevuld
                if row["Handmatige Prijs"] > 0:
                    return row["Handmatige Prijs"]
                
                # Logica voor SAP Prijs
                elif prijsbepaling_optie == "SAP prijs":
                    return row["SAP Prijs"]
                
                # Logica voor RSP
                elif prijsbepaling_optie == "RSP":
                    rsp_met_kwaliteit = row["RSP"] * (row["Prijskwaliteit"] / 100)
                    return (rsp_met_kwaliteit * 20 // 1 + (1 if (rsp_met_kwaliteit * 20 % 1) > 0 else 0)) / 20
                
                # Logica voor PricePilot
                elif prijsbepaling_optie == "PricePilot logica":
                    # Zorg ervoor dat zowel SAP Prijs als RSP niet 0 zijn
                    if row["SAP Prijs"] > 0 and row["RSP"] > 0:
                        return min(row["SAP Prijs"], row["RSP"])
                    elif row["SAP Prijs"] > 0:
                        return row["SAP Prijs"]  # Gebruik SAP Prijs als RSP 0 is
                    elif row["RSP"] > 0:
                        return row["RSP"]  # Gebruik RSP als SAP Prijs 0 is
                    else:
                        return 0  # Als beide 0 zijn, zet Prijs_backend op 0
                
                # Default naar 0 als niets anders van toepassing is
                return 0

            # Pas de prijsbepaling logica toe op de DataFrame
            df["Prijs_backend"] = df.apply(bepaal_prijs_backend, axis=1)

            # Verkoopprijs is gelijk aan Prijs_backend
            df["Verkoopprijs"] = df["Prijs_backend"]

        except Exception as e:
            st.error(f"Fout bij het berekenen van Prijs_backend: {e}")

        return df








# Controleer en zet kolommen om
for col in ["M2 totaal", "RSP", "Verkoopprijs"]:
    if col not in st.session_state.offer_df.columns:
        st.session_state.offer_df[col] = 0
    st.session_state.offer_df[col] = pd.to_numeric(st.session_state.offer_df[col], errors='coerce').fillna(0)

# Berekeningen uitvoeren
totaal_m2 = st.session_state.offer_df["M2 totaal"].sum()
totaal_bedrag = (st.session_state.offer_df["M2 totaal"] * st.session_state.offer_df["Prijs_backend"]).sum()

# Maak drie kolommen
col1, col2, col3 = st.sidebar.columns(3)

# HTML weergeven in de zijbalk
with col2:
    st.image("BullsAI_logo.png", width=int(30 / 100 * 1024))  # Pas grootte aan (30% van origineel)
    st.sidebar.markdown("---")  # Scheidingslijn voor duidelijkheid  


with tab3:    
    cutoff_value = st.slider(
        "Matchwaarde AI",
        min_value=0.01,
        max_value=1.0,
        value=0.6,  # Standaardwaarde
        step=0.01,  # Stappen in float
        help="Stel matchwaarde in. Hogere waarde betekent strengere matching, 0.6 aanbevolen."
    )
    
    # Bijlagen in mail definiëren
    def detect_relevant_columns(df):
        """
        Detecteert de relevante kolommen (Artikelnaam, Hoogte, Breedte, Aantal) in een DataFrame.
        """
        # Standaardiseer kolomnamen in de DataFrame (trim en lower)
        standardized_columns = {col: col.strip().lower() for col in df.columns}
        
        column_mapping = {
            "Artikelnaam": ["artikelnaam", "artikel", "product", "type", "article", "samenstelling", "glastype", "omschrijving"],
            "Hoogte": ["hoogte", "height", "h"],
            "Breedte": ["breedte", "width", "b"],
            "Aantal": ["aantal", "quantity", "qty", "stuks"]
        }
        detected_columns = {}
    

    
        for key, patterns in column_mapping.items():
            for pattern in patterns:
                match = [original_col for original_col, std_col in standardized_columns.items() if std_col == pattern]
                if match:
                    detected_columns[key] = match[0]
                    break
    
        return detected_columns
    
    
    

        
    # Dynamisch zoeken in de zijbalk
    with st.sidebar:
        search_query = ""
    
        # Filter de resultaten op basis van de invoer
        if not accounts_df.empty and search_query:
            filtered_df = accounts_df[accounts_df["Klantnaam"].str.contains(search_query, case=False, na=False)]
        else:
            filtered_df = accounts_df
    
        # Voeg een lege string toe als eerste optie in de lijst
        klantopties = [""] + filtered_df["Klantinfo"].tolist()
    
        # Toon de selectbox met de lege regel als standaardwaarde
        selected_customer = st.selectbox(
            "Selecteer een klant",
            options=klantopties,
            index=0,  # Hiermee wordt de lege regel standaard geselecteerd
            help="Kies een klant uit de lijst.",
        )
    
        # Afleiden van customer_number als de selectie is gemaakt
        if selected_customer:
            customer_number = selected_customer[:6]  # Haal de eerste 6 tekens uit de selectie
        else:
            customer_number = None
    
        st.session_state.customer_number = str(customer_number) if customer_number else ''
    
        # Klantreferentie invoer
        customer_reference = st.text_input(
            "Klantreferentie",
            value=st.session_state.get("customer_reference", ""),
        )



    offer_amount = totaal_bedrag

# Maak twee kolommen in de sidebar (verhoudingen kunnen aangepast worden)
col1, col2 = st.sidebar.columns([1, 1])

# Linker kolom: Totaalwaarden
with col1:
    st.metric("Totaal m2", f"{totaal_m2:.2f}")
    st.metric("Totaal Bedrag", f"€ {totaal_bedrag:.2f}")

# Rechter kolom: Klantinformatie
with col2:
    if customer_number in customer_data:
        # Haal klantinformatie op
        omzet_klant = customer_data[customer_number]['revenue'].replace("euro", "€").strip()
        klantgrootte = customer_data[customer_number]['size']
        
        # Haal de aangepaste matrix op
        prijsscherpte_matrix = st.session_state.prijsscherpte_matrix
        prijsscherpte = ""

        # Bepaal prijsscherpte op basis van klantgrootte en offertebedrag
        if klantgrootte in prijsscherpte_matrix.columns:
            for index, row in prijsscherpte_matrix.iterrows():
                if offer_amount >= row["Offertebedrag"]:
                    prijsscherpte = row[klantgrootte]
                else:
                    break
        
        # Toon klantinformatie als metrics
        st.metric("Omzet klant", omzet_klant)
        st.metric("Klantgrootte", klantgrootte)
        # st.metric("Prijsscherpte", prijsscherpte) == prijsscherpte nu eruit gehaald, maar kan getoond worden.

# Functie om synoniemen te vervangen in invoertekst
def replace_synonyms(input_text, synonyms):
    for term, synonym in synonyms.items():
        input_text = input_text.replace(term, synonym)
    return input_text

# Functie om artikelgegevens te vinden
def find_article_details(article_number):
    # Sla het originele artikelnummer op
    original_article_number = article_number  

        
    # 1. Controleer of artikelnummer een exacte match is in synonym_dict.values()
    if article_number in synonym_dict.values():
        filtered_articles = article_table[article_table['Material'].astype(str) == str(article_number)]
        if not filtered_articles.empty:
            return (
                filtered_articles.iloc[0]['Description'],
                filtered_articles.iloc[0]['Min_prijs'],
                filtered_articles.iloc[0]['Max_prijs'],
                article_number,
                "synoniem",  # Bron: exacte match in synonym_dict.values()
                article_number,  # Original article number
                None  # Fuzzy match remains empty
            )

    # 2. Controleer of artikelnummer een exacte match is in synonym_dict.keys()
    if article_number in synonym_dict.keys():
        matched_article_number = synonym_dict[article_number]  # Haal het bijbehorende artikelnummer op
        filtered_articles = article_table[article_table['Material'].astype(str) == str(matched_article_number)]
        if not filtered_articles.empty:
            return (
                filtered_articles.iloc[0]['Description'],
                filtered_articles.iloc[0]['Min_prijs'],
                filtered_articles.iloc[0]['Max_prijs'],
                matched_article_number,
                "synoniem",  # Bron: exacte match in synonym_dict.keys()
                article_number,  # Original article number
                None  # Fuzzy match remains empty
            )

    # 3. Zoek naar een bijna-match met RapidFuzz
    closest_match = process.extractOne(article_number, synonym_dict.keys(), scorer=fuzz.ratio, score_cutoff=cutoff_value * 100)
    if closest_match:
        best_match = closest_match[0]
        matched_article_number = synonym_dict[best_match]
        filtered_articles = article_table[article_table['Material'].astype(str) == str(matched_article_number)]
        if not filtered_articles.empty:
            return (
                filtered_articles.iloc[0]['Description'],
                filtered_articles.iloc[0]['Min_prijs'],
                filtered_articles.iloc[0]['Max_prijs'],
                matched_article_number,
                "interpretatie",  # Bron: RapidFuzz match
                article_number,  # Original article number
                best_match  # Fuzzy match found
            )
    
    # 4. Zoek naar een bijna-match met difflib
    closest_matches = difflib.get_close_matches(article_number, synonym_dict.keys(), n=1, cutoff=cutoff_value)
    if closest_matches:
        best_match = closest_matches[0]
        matched_article_number = synonym_dict[best_match]
        filtered_articles = article_table[article_table['Material'].astype(str) == str(matched_article_number)]
        if not filtered_articles.empty:
            return (
                filtered_articles.iloc[0]['Description'],
                filtered_articles.iloc[0]['Min_prijs'],
                filtered_articles.iloc[0]['Max_prijs'],
                matched_article_number,
                "interpretatie",  # Bron: difflib match
                article_number,  # Original article number
                best_match  # Fuzzy match found
            )

    #  # 5. Zoek alternatieven via GPT
    # synonym_list_str = "\n".join([f"{k}: {v}" for k, v in synonym_dict.items()])
    # prompt = f"""
    # Op basis van voorgaande regex is de input '{original_article_number}' niet toegewezen aan een synoniem. Hier is een lijst van beschikbare synoniemen:
    # {synonym_list_str}
    # Kun je één synoniem voorstellen die het dichtst in de buurt komt bij '{original_article_number}'? Onthoud, het is enorm belangrijk dat je slechts het synoniem retourneert, geen begeleidend schrijven.
    # """
    # try:
    #     response = openai.chat.completions.create(
    #         model="gpt-3.5-turbo",
    #         messages=[
    #             {"role": "system", "content": "Je bent een behulpzame assistent die een synoniem zoekt dat het dichtst in de buurt komt van het gegeven artikelnummer. Het is enorm belangrijk dat je slechts het synoniem retourneert, geen begeleidend schrijven."},
    #             {"role": "user", "content": prompt}
    #         ],
    #         max_tokens=20, 
    #         temperature=0.5,
    #     )
    
    #     response_text = response.choices[0].message.content.strip()
    
    #     # Gebruik de GPT-response correct
    #     best_guess = response_text.split("\n")[0] if "\n" in response_text else response_text
    #     matched_article_number = synonym_dict.get(best_guess, best_guess)
    
    #     # Verifieer of het gegenereerde synoniem geldig is
    #     filtered_articles = article_table[article_table['Material'].astype(str) == str(matched_article_number)]
    #     if not filtered_articles.empty:
    #         return (
    #             filtered_articles.iloc[0]['Description'],  # Artikelnaam
    #             filtered_articles.iloc[0]['Min_prijs'],
    #             filtered_articles.iloc[0]['Max_prijs'],
    #             matched_article_number,  # Artikelnummer
    #             "GPT",  # Bron: GPT match
    #             original_article_number,  # Original article number
    #             best_guess  # Fuzzy match gevonden door GPT
    #         )
    
    # except Exception as e:
    #     st.warning(f"Fout bij het raadplegen van OpenAI API: {e}")



    # 6. Als alles niet matcht
    return (article_number, None, None, original_article_number, "niet gevonden", original_article_number, None)


# Werkt de artikelnummer bij in de DataFrame op basis van de ingevulde artikelnaam. Gebruikt fuzzy matching om de beste overeenkomst te vinden.
def update_article_numbers_from_names(df, article_table, cutoff_value = cutoff_value):
    if df.empty or article_table.empty:
        return df  # Return ongeldige invoer

    for index, row in df.iterrows():
        artikelnaam = row.get("Artikelnaam", "").strip()

        # Alleen bijwerken als er een naam is en de artikelnummer ontbreekt of een slechte match is
        if artikelnaam and (pd.isna(row.get("Artikelnummer")) or row["Source"] in ["niet gevonden", "GPT"]):

            # Zoek de beste match met fuzzy matching
            best_match = process.extractOne(artikelnaam, article_table["Description"], scorer=fuzz.ratio, score_cutoff=cutoff_value * 100)

            if best_match:
                best_article_name, score, match_index = best_match
                matched_article_number = article_table.iloc[match_index]["Material"]

                df.at[index, "Artikelnummer"] = matched_article_number
                df.at[index, "Source"] = "fuzzy_match"  # Markeer als fuzzy match
                df.at[index, "fuzzy_match"] = best_article_name  # Voeg fuzzy match kolom toe
            else:
                df.at[index, "Source"] = "niet gevonden"  # Geen match gevonden

    return df



# Functie om aanbevolen prijs te berekenen
def calculate_recommended_price(min_price, max_price, prijsscherpte):
    if min_price is not None and max_price is not None and prijsscherpte != "":
        return min_price + ((max_price - min_price) * (100 - prijsscherpte) / 100)
    return None


# Functie om m2 per stuk te berekenen
def calculate_m2_per_piece(width, height):
    if width and height:
        width_m = int(width) / 1000
        height_m = int(height) / 1000
        m2 = max(width_m * height_m, 0.65)
        return m2
    return None

# Functie om determine_spacer waarde te bepalen uit samenstellingstekst
def determine_spacer(term, default_value="15 - alu"):
    if term and isinstance(term, str):

        # Regex: Zoek alleen naar getallen tussen `-` en `-`
        matches = re.findall(r'-(\d+)-', term)
        values = list(map(int, matches))  # Converteer de gevonden matches naar integers

        # Controleer of er minimaal één getal is gevonden
        if len(values) >= 1:
            spacer_value = values[0]  # Pak de eerste gevonden waarde

            # Controleer of de waarde binnen de juiste range ligt
            if 3 < spacer_value < 30:
                if any(keyword in term.lower() for keyword in ["we", "warmedge", "warm edge"]):
                    result = f"{spacer_value} - warm edge"
                else:
                    result = f"{spacer_value} - alu"
                return result
    return default_value


# Voorbeeld van hoe de waarde wordt opgeslagen in de state
def update_spacer_state(user_input, app_state):
    selected_spacer = determine_spacer(user_input)
    app_state["spacer"] = selected_spacer


# Functie om bestaande spacers niet te overschrijven bij updates
def preserve_existing_spacers(df):
    for index, row in df.iterrows():
        if pd.notna(row.get("Spacer")):
            continue  # Behoud bestaande waarde
        # Alleen waarden aanpassen als deze niet bestaan of leeg zijn
        df.at[index, "Spacer"] = determine_spacer(row.get("Spacer", "15 - alu"))
    return df


def update_offer_data(df):
    for index, row in df.iterrows():
        if pd.notna(row['Breedte']) and pd.notna(row['Hoogte']):
            df.at[index, 'M2 p/s'] = calculate_m2_per_piece(row['Breedte'], row['Hoogte'])
        if pd.notna(row['Aantal']) and pd.notna(df.at[index, 'M2 p/s']):
            df.at[index, 'M2 totaal'] = float(row['Aantal']) * float(str(df.at[index, 'M2 p/s']).split()[0].replace(',', '.'))
        if pd.notna(row['Artikelnummer']):
            # Controleer of Source al is gevuld
            if pd.isna(row.get('Source')) or row['Source'] in ['niet gevonden', 'GPT']:
                description, min_price, max_price, article_number, source, original_article_number, fuzzy_match = find_article_details(row['Artikelnummer'])
                if description:
                    df.at[index, 'Artikelnaam'] = description
                if min_price is not None and max_price is not None:
                    df.at[index, 'Min_prijs'] = min_price
                    df.at[index, 'Max_prijs'] = max_price
                if source:  # Alleen Source bijwerken als deze leeg is
                    df.at[index, 'Source'] = source
                if original_article_number:
                    df.at[index, 'original_article_number'] = original_article_number
                if fuzzy_match:
                    df.at[index, 'fuzzy_match'] = fuzzy_match
            
            # Update SAP Prijs
            if st.session_state.customer_number in sap_prices:
                sap_prijs = sap_prices[st.session_state.customer_number].get(row['Artikelnummer'], None)
                df.at[index, 'SAP Prijs'] = sap_prijs if sap_prijs else None
            else:
                df.at[index, 'SAP Prijs'] = None
    df = bereken_prijs_backend(df)
    return df


# Functie om de RSP voor alle regels te updaten
def update_rsp_for_all_rows(df, prijsscherpte):
    if prijsscherpte:

        def calculate_rsp(row):
            min_price = row.get('Min_prijs', None)
            max_price = row.get('Max_prijs', None)
            if pd.notna(min_price) and pd.notna(max_price):
                rsp_value = calculate_recommended_price(min_price, max_price, prijsscherpte)
                return round(rsp_value * 20) / 20
            return row.get('RSP', None)

        df['RSP'] = df.apply(calculate_rsp, axis=1)

        # Pas backend-berekeningen toe
        df = bereken_prijs_backend(df)

    return df




# Functie om Prijs_backend te updaten na wijzigingen
def update_prijs_backend():
    st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)

def reset_rijnummers(df):
    if not df.empty:
        df['Rijnummer'] = range(1, len(df) + 1)
    return df


# JavaScript-code voor conditionele opmaak
cell_style_js = JsCode("""
function(params) {
    if (params.colDef.field === "RSP" && params.data.Prijs_backend === params.data.RSP) {
        return {'backgroundColor': '#DFFFD6', 'fontWeight': 'bold'};  // Lichtgroen met vetgedrukte letters
    } else if (params.colDef.field === "SAP Prijs" && params.data.Prijs_backend === params.data["SAP Prijs"]) {
        return {'backgroundColor': '#DFFFD6', 'fontWeight': 'bold'};  // Lichtgroen met vetgedrukte letters
    } else if (params.colDef.field === "Verkoopprijs" && params.data.Prijs_backend === params.data.Verkoopprijs) {
        return {'backgroundColor': '#DFFFD6', 'fontWeight': 'bold'};  // Lichtgroen met vetgedrukte letters
    } else if (params.colDef.field !== "Verkoopprijs") {
        return {'backgroundColor': '#e0e0e0'};  // Grijs voor alle andere cellen
    }
    return null;
}
""")

# Voeg een cell renderer toe om de stericoon weer te geven
cell_renderer_js = JsCode("""
function(params) {
    if (params.data.Source === "interpretatie" || params.data.Source === "GPT") {
        return `✨ ${params.value}`;  // Voeg stericoon toe vóór de waarde
    }
    return params.value;  // Toon de originele waarde
}
""")


def save_changes(df):
    st.session_state.offer_df = df
    st.session_state.offer_df = update_offer_data(st.session_state.offer_df)
    st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)
    st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, st.session_state.get('prijsscherpte', ''))

# Offerte Genereren tab
with tab1:
    # Voeg een veld toe voor prijskwaliteit als RSP wordt gekozen met beperkte breedte
    if prijsbepaling_optie == "RSP":
        col1, _ = st.columns([1, 10])
        with col1:
            prijskwaliteit = st.number_input("Prijskwaliteit (%)", min_value=0, max_value=200, value=100, key="prijskwaliteit")
        st.session_state.offer_df["Prijskwaliteit"] = prijskwaliteit

    # Altijd de logica via de functie bereken_prijs_backend toepassen
    st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)

# JavaScript code voor het opslaan van wijzigingen
js_update_code = JsCode('''
function onCellEditingStopped(params) {
    // Opslaan van gewijzigde data na het bewerken van een cel
    let updatedRow = params.node.data;

    // Zorg ervoor dat wijzigingen worden doorgevoerd in de grid
    params.api.applyTransaction({ update: [updatedRow] });
}
''')


# Maak grid-opties aan voor AgGrid met gebruik van een "select all" checkbox in de header
gb = GridOptionsBuilder.from_dataframe(st.session_state.offer_df)
gb.configure_default_column(flex=1, minWidth=50, editable=True)
gb.configure_column("Spacer", editable=True, cellEditor='agSelectCellEditor', cellEditorParams={"values": ["4 - alu", "6 - alu", "7 - alu", "8 - alu", "9 - alu", "10 - alu", "12 - alu", "13 - alu", "14 - alu", "15 - alu", "16 - alu", "18 - alu", "20 - alu", "24 - alu", "10 - warm edge", "12 - warm edge", "14 - warm edge", "15 - warm edge", "16 - warm edge", "18 - warm edge", "20 - warm edge", "24 - warm edge"]})
gb.configure_column("Rijnummer", type=["numericColumn"], editable=False, cellStyle={"backgroundColor": "#e0e0e0"}, cellRenderer=cell_renderer_js)
gb.configure_column("Artikelnaam", width=600)
gb.configure_column("Offertenummer", hide=True)
gb.configure_column("Prijs_backend", hide=True)
gb.configure_column("Min_prijs", hide=True)
gb.configure_column("Artikelnummer", hide=False)
gb.configure_column("Prijskwaliteit", hide=True)
gb.configure_column("Max_prijs", hide=True)
gb.configure_column("Handmatige Prijs", editable=True, type=["numericColumn"])
gb.configure_column("Breedte", editable=True, type=["numericColumn"])
gb.configure_column("Hoogte", editable=True, type=["numericColumn"])
gb.configure_column("Aantal", editable=True, type=["numericColumn"])
gb.configure_column("RSP", editable=False, type=["numericColumn"], valueFormatter="x.toFixed(2)", cellStyle=cell_style_js)
gb.configure_column("Verkoopprijs", editable=True, type=["numericColumn"], cellStyle=cell_style_js, valueFormatter="x.toFixed(2)")
gb.configure_column("M2 p/s", editable=False, type=["numericColumn"], cellStyle={"backgroundColor": "#e0e0e0"}, valueFormatter="x.toFixed(2)")
gb.configure_column("M2 totaal", editable=False, type=["numericColumn"], cellStyle={"backgroundColor": "#e0e0e0"}, valueFormatter="x.toFixed(2)")
gb.configure_column("SAP Prijs", editable=False, type=["numericColumn"], valueFormatter="x.toFixed(2)", cellStyle=cell_style_js)
gb.configure_column("Source", hide=True)


# Configuratie voor selectie, inclusief checkbox in de header voor "select all"
gb.configure_selection(
    selection_mode='multiple',
    use_checkbox=True,
    header_checkbox=True  # Voeg een selectievakje in de header toe
)

# Voeg de JavaScript code toe aan de grid-opties
gb.configure_grid_options(onCellEditingStopped=js_update_code)

# Overige configuratie van de grid
gb.configure_grid_options(domLayout='normal', rowHeight=23)  # Dit zorgt ervoor dat scrollen mogelijk is

# Voeg een JavaScript event listener toe voor updates bij het indrukken van Enter
js_update_code = JsCode('''
function onCellValueChanged(params) {
    let rowNode = params.node;
    let data = rowNode.data;

    // Zorg ervoor dat wijzigingen direct worden toegepast
    params.api.applyTransaction({ update: [data] });

    // Forceer visuele update
    params.api.refreshCells({ force: true });

    // Luister naar de Enter-toets
    document.addEventListener('keydown', function(event) {
        if (event.key === 'Enter') {
            // Ververs de grid wanneer Enter wordt ingedrukt
            params.api.redrawRows();
        }
    });
}
''')
gb.configure_grid_options(onCellValueChanged=js_update_code)

# Bouw grid-opties
grid_options = gb.build()

# Offerte Genereren tab
with tab1:

    # Toon de AG Grid met het material-thema
    edited_df_response = AgGrid(
        st.session_state.offer_df,
        gridOptions=grid_options,
        theme='alpine',
        fit_columns_on_grid_load=True,
        enable_enterprise_modules=True,
        update_mode=GridUpdateMode.SELECTION_CHANGED,
        columns_auto_size_mode=ColumnsAutoSizeMode.FIT_CONTENTS,
        allow_unsafe_jscode=True
    )

    # Update de DataFrame na elke wijziging
    if "data" in edited_df_response:
        updated_df = pd.DataFrame(edited_df_response['data'])
        # Werk de sessiestatus bij met de nieuwe data
        st.session_state.offer_df = updated_df
        # Voer alle benodigde berekeningen uit
        st.session_state.offer_df = update_offer_data(st.session_state.offer_df)
        st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)

 
    # Verbeterde update_tabel functie
def update_tabel():
    updated_df = pd.DataFrame(edited_df_response['data'])
    st.session_state.offer_df = updated_df
    st.session_state.offer_df = update_offer_data(st.session_state.offer_df)
    st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)

    new_df = st.session_state.offer_df
      
    st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_df], ignore_index=True)
    st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, prijsscherpte)
    st.session_state["trigger_update"] = True
    st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)

# Offerte Genereren tab
with tab1:
    
    # Knop om de tabel bij te werken
    if st.button("Update tabel"):
        update_tabel()


 
    # Update de DataFrame na elke wijziging
    updated_df = edited_df_response['data']
    save_changes(pd.DataFrame(updated_df))
    
    # Sla de geselecteerde rijen op in sessie status
    selected_rows = edited_df_response.get('selected_rows_id', edited_df_response.get('selected_rows', edited_df_response.get('selected_data', [])))


    # Zorg dat selected_rows geen None of DataFrame is, maar altijd een lijst
    if selected_rows is None or not isinstance(selected_rows, list):
        selected_rows = []
    
    # Als er rijen zijn geselecteerd, zet deze in de sessie state
    if isinstance(selected_rows, list) and len(selected_rows) > 0:
        try:
            st.session_state.selected_rows = [int(r) for r in selected_rows]
        except ValueError:
            st.write("Waarschuwing: Fout bij het converteren van geselecteerde rijen naar indices.")
    else:
        st.session_state.selected_rows = []
    
    def delete_selected_rows(df, selected):
        if selected_rows is not None and len(selected_rows) > 0:
            # Zorg ervoor dat de indices integers zijn
            selected = [int(i) for i in selected]
            st.write("Geselecteerde indices na conversie:", selected)  # Debugging statement
    
            # Verwijder de geselecteerde rijen en reset de index
            new_df = df.drop(index=selected_rows, errors='ignore').reset_index(drop=True)
            return new_df
           
        else:
            return df

   
    # Knoppen toevoegen aan de GUI
    col1, col2, col3, col4, col5, col6 = st.columns(6)
    with col1:
        if st.button("Voeg rij toe"):
            # Voeg een lege rij toe aan het DataFrame
            new_row = pd.DataFrame({
                "Offertenummer": [None], "Artikelnaam": [""], "Artikelnummer": [""], "Spacer": ["15 - alu"], "Breedte": [0], "Hoogte": [0],
                "Aantal": [0], "RSP": [None], "M2 p/s": [0], "M2 totaal": [0], "Min_prijs": [None], "Max_prijs": [None], "Handmatige Prijs": [1000]
            })
            st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_row], ignore_index=True)
            st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)
            # Werk de Rijnummer-kolom bij zodat deze overeenkomt met de index + 1
            st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
            # Vernieuw de AgGrid
            st.rerun()
    
    with col2:
        if st.button("Verwijder rijen (2x klikken)", key='delete_rows_button'):
            # Haal de geselecteerde rijen op in de juiste vorm
            selected = st.session_state.selected_rows
            st.write("Geselecteerde rijen voor verwijdering:", selected)  # Debugging statement
    
            # Verwijder rijen op basis van index
            if len(selected) > 0:
                # Verwijder de rijen uit de DataFrame op basis van de geselecteerde indices
                st.session_state.offer_df = delete_selected_rows(st.session_state.offer_df, selected)
                st.session_state.selected_rows = []  # Reset de geselecteerde rijen na verwijderen
                # Reset de Rijnummer-kolom na verwijderen
                st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
                st.rerun()
            else:
                st.warning("Selecteer eerst rijen om te verwijderen.")
    
        # Zorg dat de update wordt getriggerd na verwijdering
        st.session_state['trigger_update'] = True

  
# Functie om getallen van 1 tot 100 te herkennen
def extract_numbers(text):
    pattern = r'\b(1|[1-9]|[1-9][0-9]|100)\b'
    matches = re.findall(pattern, text)
    return [int(match) for match in matches]

# Functie om woorden naar getallen om te zetten
def word_to_number(word):
    mapping = {
        "een": 1, "twee": 2, "drie": 3, "vier": 4, "vijf": 5, "zes": 6, "zeven": 7, "acht": 8, "negen": 9, "tien": 10, "elf": 11, "twaalf": 12, "dertien": 13, "veertien": 14, "vijftien": 15, "zestien": 16, "zeventien": 17, "achttien": 18, 
        "negentien": 19, "twintig": 20, "eenentwintig": 21, "tweeëntwintig": 22, "drieëntwintig": 23, "vierentwintig": 24, "vijfentwintig": 25, "zesentwintig": 26, "zevenentwintig": 27, "achtentwintig": 28, 
        "negenentwintig": 29, "dertig": 30, "eenendertig": 31, "tweeëndertig": 32, "drieënendertig": 33, "vierendertig": 34, "vijfendertig": 35, "zesendertig": 36, "zevenendertig": 37, "achtendertig": 38, 
        "negenendertig": 39, "veertig": 40, "eenenveertig": 41, "tweeënveertig": 42, "drieënveertig": 43, "vierenveertig": 44, "vijfenveertig": 45, "zesenveertig": 46, "zevenenveertig": 47, "achtenveertig": 48, 
        "negenenveertig": 49, "vijftig": 50, "eenenvijftig": 51, "tweeënvijftig": 52, "drieënvijftig": 53, "vierenvijftig": 54, "vijfenvijftig": 55, "zesenvijftig": 56, "zevenenvijftig": 57, "achtenvijftig": 58, 
        "negenenvijftig": 59, "zestig": 60, "eenenzestig": 61, "tweeënzestig": 62, "drieënzestig": 63, "vierenzestig": 64, "vijfenzestig": 65, "zesenzestig": 66, "zevenenzestig": 67, "achtenzestig": 68, 
        "negenenzestig": 69, "zeventig": 70, "eenenzeventig": 71, "tweeënzeventig": 72, "drieënzeventig": 73, "vierenzeventig": 74, "vijfenzeventig": 75, "zesenzeventig": 76, "zevenenzeventig": 77, "achtenzeventig": 78, 
        "negenenzeventig": 79, "tachtig": 80, "eenentachtig": 81, "tweeëntachtig": 82, "drieëntachtig": 83, "vierentachtig": 84, "vijfentachtig": 85, "zesentachtig": 86, "zevenentachtig": 87, "achtentachtig": 88, 
        "negenentachtig": 89, "negentig": 90, "eenennegentig": 91, "tweeënnegentig": 92, "drieënnegentig": 93, "vierennegentig": 94, "vijfennegentig": 95, "zesennegentig": 96, "zevenennegentig": 97, "achtennegentig": 98, 
        "negenennegentig": 99, "honderd": 100
    }
    return mapping.get(word, None)

# Callback functie voor het verwijderen van geselecteerde rijen
@st.cache_data
def update_dash_table(n_dlt, n_add, data):
    if ctx.triggered_id == "add-row-btn":
        new_row = pd.DataFrame({
            "Offertenummer": [None],
            "Artikelnaam": [""],
            "Artikelnummer": [""],
            "Spacer": [st.session_state.get("last_selected_spacer", "15 - alu")],  # Gebruik de laatst geselecteerde waarde
            "Breedte": [0],
            "Hoogte": [0],
            "Aantal": [0],
            "RSP": [0],
            "M2 p/s": [0],
            "M2 totaal": [0],
            "Min_prijs": [0],
            "Max_prijs": [0],
            "Verkoopprijs": [0]
        })
        df_new_row = pd.DataFrame(new_row)
        updated_table = pd.concat([pd.DataFrame(data), df_new_row])
        return False, updated_table.to_dict("records")

    elif ctx.triggered_id == "delete-row-btn":
        return True, no_update


  
# Functie om het aantal uit tekst te extraheren
def extract_quantity(text):
    # Zoek naar een getal of woord dat voor 'stuks', 'aantal', 'ruiten', 'st', 'keer', of 'x' staat
    unit_matches = re.findall(r'(\d+|twee|drie|vier|vijf|zes|zeven|acht|negen|tien|elf|twaalf|dertien|veertien|vijftien|zestien|zeventien|achttien|negentien|twintig|eenentwintig|tweeëntwintig|drieëntwintig|vierentwintig|vijfentwintig|zesentwintig|zevenentwintig|achtentwintig|negenentwintig|dertig|eenendertig|tweeëndertig|drieëndertig|vierendertig|vijfendertig|zesendertig|zevenendertig|achtendertig|negenendertig|veertig|eenenveertig|tweeënveertig|drieënveertig|vierenveertig|vijfenveertig|zesenveertig|zevenenveertig|achtenveertig|negenenveertig|vijftig|eenenvijftig|tweeënvijftig|drieënvijftig|vierenvijftig|vijfenvijftig|zesenvijftig|zevenenvijftig|achtenvijftig|negenenvijftig|zestig|eenenzestig|tweeënzestig|drieënzestig|vierenzestig|vijfenzestig|zesenzestig|zevenenzestig|achtenzestig|negenenzestig|zeventig|eenenzeventig|tweeënzeventig|drieënzeventig|vierenzeventig|vijfenzeventig|zesenzeventig|zevenenzeventig|achtenzeventig|negenenzeventig|tachtig|eenentachtig|tweeëntachtig|drieëntachtig|vierentachtig|vijfentachtig|zesentachtig|zevenentachtig|achtentachtig|negenentachtig|negentig|eenennegentig|tweeënnegentig|drieënnegentig|vierennegentig|vijfennegentig|zesennegentig|zevenennegentig|achtennegentig|negenennegentig|honderd)\s*(stuks|aantal|ruiten|st|keer|x)\b', text, re.IGNORECASE)

    
    if unit_matches:
        # Als een match gevonden is, zet het om naar een getal
        return word_to_number(unit_matches[0][0]) if unit_matches[0][0].isalpha() else int(unit_matches[0][0])
    
    # Anders zoek naar een getal alleen
    quantity_matches = extract_numbers(text)
    word_matches = re.findall(r'\b(twee|drie|vier|vijf|zes|zeven|acht|negen|tien|elf|twaalf|dertien|veertien|vijftien|zestien|zeventien|achttien|negentien|twintig|eenentwintig|tweeëntwintig|drieëntwintig|vierentwintig|vijfentwintig|zesentwintig|zevenentwintig|achtentwintig|negenentwintig|dertig|eenendertig|tweeëndertig|drieëndertig|vierendertig|vijfendertig|zesendertig|zevenendertig|achtendertig|negenendertig|veertig|eenenveertig|tweeënveertig|drieënveertig|vierenveertig|vijfenveertig|zesenveertig|zevenenveertig|achtenveertig|negenenveertig|vijftig|eenenvijftig|tweeënvijftig|drieënvijftig|vierenvijftig|vijfenvijftig|zesenvijftig|zevenenvijftig|achtenvijftig|negenenvijftig|zestig|eenenzestig|tweeënzestig|drieënzestig|vierenzestig|vijfenzestig|zesenzestig|zevenenzestig|achtenzestig|negenenzestig|zeventig|eenenzeventig|tweeënzeventig|drieënzeventig|vierenzeventig|vijfenzeventig|zesenzeventig|zevenenzeventig|achtenzeventig|negenenzeventig|tachtig|eenentachtig|tweeëntachtig|drieëntachtig|vierentachtig|vijfentachtig|zesentachtig|zevenentachtig|achtentachtig|negenentachtig|negentig|eenennegentig|tweeënnegentig|drieënnegentig|vierennegentig|vijfennegentig|zesennegentig|zevenennegentig|achtennegentig|negenennegentig|honderd)\b', text)

    if word_matches:
        return word_to_number(word_matches[0])  # Neem het eerste gevonden aantal in woorden
    if quantity_matches:
        return quantity_matches[0]  # Neem het eerste gevonden aantal in cijfers
    return None


# Functie om afmetingen (breedte en hoogte) uit tekst te extraheren
def extract_dimensions(text):
    """
    Extraheert breedte en hoogte uit een regel tekst.
    """
    matches = re.findall(r'(\d{3,4})\s*[xX*]\s*(\d{3,4})', text)  # Herken 800x900 of 800*900
    if matches:
        return int(matches[0][0]), int(matches[0][1])  # Eerste geldige combinatie teruggeven
    
    # Alternatieve notatie zoals "700 bij 800"
    matches = re.findall(r'(\d{2,4})\s*bij\s*(\d{2,4})', text, re.IGNORECASE)
    if matches:
        return int(matches[0][0]), int(matches[0][1])

    return None, None

# Functie om alle gegevens (aantal, afmetingen, artikelnummer) te extraheren
def extract_all_details(line):
    # Extract quantity
    quantity = extract_quantity(line)
    # Extract dimensions
    width, height = extract_dimensions(line)
    # Extract article number
    article_number_match = re.search(r'(\d+[./-]?\d*[-*#]\d+[./-]?\d*)', line)
    article_number = article_number_match.group(0) if article_number_match else None
    return quantity, width, height, article_number

def handle_gpt_chat():
    if customer_input:
        lines = customer_input.splitlines()
        data = []
        current_article_number = None  # Huidig artikelnummer onthouden
        
        for line in lines:
            # Controleer of er een artikelnummer in de regel staat
            detected_article_number = re.search(r'(\d+[./-]?\d*[-*#]\d+[./-]?\d*)', line)
            if detected_article_number:
                current_article_number = detected_article_number.group(0)  # Update huidig artikelnummer


            # Probeer m2-formaat en artikelnummer te detecteren
            m2_match = re.search(r'(\d+)\s*m2.*?(\d+-\d+)|^(\d+-\d+).*?(\d+)\s*m2', line, re.IGNORECASE)

            # Extract details zoals aantal, breedte en hoogte
            quantity, width, height, article_number = extract_all_details(line)

            # Als er geen artikelnummer in deze regel staat, gebruik de vorige (indien beschikbaar)
            if not article_number and current_article_number:
                article_number = current_article_number
                st.sidebar.info(f"Geen nieuw artikelnummer gevonden, gebruik vorige: {article_number}")

            # Verwerking als er een m2-match is
            if m2_match:
                if m2_match.group(1):
                    m2_total = int(m2_match.group(1))
                    article_number = m2_match.group(2)
                else:
                    article_number = m2_match.group(3)
                    m2_total = int(m2_match.group(4))

                # Gebruik het synoniemenwoordenboek
                article_number = synonym_dict.get(article_number, article_number)

                # Zoek artikelgegevens op
                description, min_price, max_price, article_number, source, original_article_number, fuzzy_match = find_article_details(article_number)

                if description:
                    recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
                    verkoopprijs = None
                    prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

                    data.append([
                        None,
                        description,
                        article_number,
                        None,
                        None,
                        None,
                        None,
                        None,
                        f"{m2_total:.2f}" if m2_total is not None else None,
                        f"{recommended_price:.2f}" if recommended_price is not None else 0,
                        None,
                        None,
                        min_price,
                        max_price,
                        verkoopprijs,
                        prijs_backend,
                        source,
                        fuzzy_match,
                        original_article_number
                    ])
                else:
                    st.sidebar.warning(f"Artikelnummer '{article_number}' niet gevonden in de artikelentabel.")

            # Verwerking als er een aantal + breedte/hoogte is
            elif quantity and (width and height):
                # Zoek artikelgegevens op
                article_number = synonym_dict.get(article_number, article_number)
                description, min_price, max_price, article_number, source, original_article_number, fuzzy_match = find_article_details(article_number)

                if description:
                    spacer = determine_spacer(line)
                    m2_per_piece = round(calculate_m2_per_piece(width, height), 2)
                    m2_total = round(float(quantity) * m2_per_piece, 2)

                    recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
                    verkoopprijs = None
                    prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

                    data.append([
                        None,
                        description,
                        article_number,
                        spacer,
                        width,
                        height,
                        quantity,
                        f"{m2_per_piece:.2f}" if m2_per_piece is not None else None,
                        f"{m2_total:.2f}" if m2_total is not None else None,
                        f"{recommended_price:.2f}" if recommended_price is not None else 0,
                        None,
                        None,
                        min_price,
                        max_price,
                        verkoopprijs,
                        prijs_backend,
                        source,
                        fuzzy_match,
                        original_article_number
                    ])
                else:
                    st.sidebar.warning(f"Artikelnummer '{article_number}' niet gevonden in de artikelentabel.")
            else:
                st.sidebar.warning("Regel genegeerd: geen geldige breedte, hoogte of aantal gevonden.")

        # Als data is verzameld, voeg het toe aan de offerte-overzichtstabel
        if data:
            new_df = pd.DataFrame(data, columns=["Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", "Aantal", "M2 p/s", "M2 totaal", "RSP", "SAP Prijs", "Handmatige Prijs", "Min_prijs", "Max_prijs", "Verkoopprijs", "Prijs_backend", "Source", "fuzzy_match", "original_article_number"])
            new_df.insert(0, 'Rijnummer', new_df.index + 1)

            st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_df], ignore_index=True)
            st.session_state.offer_df = update_offer_data(st.session_state.offer_df)
            st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, prijsscherpte)
            st.session_state["trigger_update"] = True
            st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
            st.session_state.offer_df = update_article_numbers_from_names(st.session_state.offer_df, article_table)
            st.rerun()

           

        else:
            st.sidebar.warning("Geen gegevens gevonden om toe te voegen.")
    elif customer_file:
        handle_file_upload(customer_file)
    else:
        st.sidebar.warning("Voer alstublieft tekst in of upload een bestand.")

#  Werkt de SAP Prijs bij op basis van het klantnummer en artikelnummer.
def update_sap_prices(df):   
    for index, row in df.iterrows():
        artikelnummer = row.get('Artikelnummer')
        if artikelnummer and st.session_state.customer_number in sap_prices:
            df.at[index, 'SAP Prijs'] = sap_prices[st.session_state.customer_number].get(artikelnummer, None)
        else:
            df.at[index, 'SAP Prijs'] = None
    df = bereken_prijs_backend(df)  # Herbereken de prijzen
    return df  # Nu staat return binnen de functie
    
# Functie direct uitvoeren en opslaan in sessiestatus
st.session_state.offer_df = update_sap_prices(st.session_state.offer_df)

# Functie voor het verwerken van e-mailinhoud naar offerte
def handle_email_to_offer(email_body):
    if email_body:
        lines = email_body.splitlines()
        data = []
        current_article_number = None  # Huidig artikelnummer onthouden
        
        for line in lines:
            # Controleer of er een artikelnummer in de regel staat
            detected_article_number = re.search(r'(\d+[./-]?\d*[-*#]\d+[./-]?\d*)', line)
            if detected_article_number:
                current_article_number = detected_article_number.group(0)  # Update huidig artikelnummer
                st.sidebar.info(f"Nieuw artikelnummer gevonden: {current_article_number}")

            # Probeer m²-formaat en artikelnummer te detecteren
            m2_match = re.search(r'(\d+)\s*m2.*?(\d+-\d+)|^(\d+-\d+).*?(\d+)\s*m2', line, re.IGNORECASE)

            # Extract details zoals aantal, breedte en hoogte
            quantity, width, height, article_number = extract_all_details(line)

            # Als er geen artikelnummer in deze regel staat, gebruik de vorige (indien beschikbaar)
            if not article_number and current_article_number:
                article_number = current_article_number
                st.sidebar.info(f"Geen nieuw artikelnummer gevonden, gebruik vorige: {article_number}")

            # Verwerking als er een m²-match is
            if m2_match:
                if m2_match.group(1):
                    m2_total = int(m2_match.group(1))
                    article_number = m2_match.group(2)
                else:
                    article_number = m2_match.group(3)
                    m2_total = int(m2_match.group(4))

                # Gebruik het synoniemenwoordenboek
                article_number = synonym_dict.get(article_number, article_number)

                # Zoek artikelgegevens op
                description, min_price, max_price, article_number, source, original_article_number, fuzzy_match = find_article_details(article_number)

                if description:
                    recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
                    verkoopprijs = None
                    prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

                    data.append([
                        None, description, article_number, None, None, None, None,
                        None, f"{m2_total:.2f}",
                        f"{recommended_price:.2f}" if recommended_price is not None else 0,
                        None, None, min_price, max_price, verkoopprijs, prijs_backend,
                        source, fuzzy_match, original_article_number
                    ])
                else:
                    st.sidebar.warning(f"Artikelnummer '{article_number}' niet gevonden in de artikelentabel.")

            # Verwerking als er een aantal + breedte/hoogte is
            elif quantity and (width and height):
                # Zoek artikelgegevens op
                article_number = synonym_dict.get(article_number, article_number)
                description, min_price, max_price, article_number, source, original_article_number, fuzzy_match = find_article_details(article_number)

                if description:
                    spacer = determine_spacer(line)
                    m2_per_piece = round(calculate_m2_per_piece(width, height), 2)
                    m2_total = round(float(quantity) * m2_per_piece, 2)

                    recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
                    verkoopprijs = None
                    prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

                    data.append([
                        None, description, article_number, spacer, width, height, quantity,
                        f"{m2_per_piece:.2f}" if m2_per_piece is not None else None,
                        f"{m2_total:.2f}" if m2_total is not None else None,
                        f"{recommended_price:.2f}" if recommended_price is not None else 0,
                        min_price, None, max_price, None, verkoopprijs, prijs_backend,
                        source, fuzzy_match, original_article_number
                    ])
                else:
                    st.sidebar.warning(f"Artikelnummer '{article_number}' niet gevonden in de artikelentabel.")
            else:
                st.sidebar.warning("Regel genegeerd: geen geldige breedte, hoogte of aantal gevonden.")

        # Als data is verzameld, voeg het toe aan de offerte-overzichtstabel
        if data:
            new_df = pd.DataFrame(data, columns=[
                "Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", 
                "Aantal", "M2 p/s", "M2 totaal", "RSP", "SAP Prijs", "Handmatige Prijs", 
                "Min_prijs", "Max_prijs", "Verkoopprijs", "Prijs_backend", "Source", "fuzzy_match", "original_article_number"
            ])
            st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_df], ignore_index=True)
            st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, prijsscherpte)
            st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
            st.rerun()
        else:
            st.sidebar.warning("Geen gegevens gevonden om toe te voegen.")





def handle_mapped_data_to_offer(df):
    """
    Verwerkt de gemapte data en vertaalt deze naar de tabelstructuur voor offertes.
    """
    data = []
    for _, row in df.iterrows():
        description = row["Artikelnaam"]
        height = row["Hoogte"]
        width = row["Breedte"]
        quantity = row["Aantal"]

        # Synoniem lookup en artikelgegevens ophalen
        article_number = synonym_dict.get(description, description)
        description, min_price, max_price, article_number, source, original_article_number, fuzzy_match = find_article_details(article_number)

        if description:
            recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
            verkoopprijs = None
            prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

            m2_per_piece = round(calculate_m2_per_piece(width, height), 2) if width and height else None
            m2_total = round(float(quantity) * m2_per_piece, 2) if m2_per_piece and quantity else None

            data.append([
                None, description, article_number, None, width, height, quantity,
                f"{m2_per_piece:.2f}" if m2_per_piece is not None else None,
                f"{m2_total:.2f}" if m2_total is not None else None,
                f"{recommended_price:.2f}" if recommended_price is not None else 0,
                min_price, None, max_price, None, verkoopprijs, prijs_backend,
                source, fuzzy_match, original_article_number
            ])
        else:
            st.sidebar.warning(f"Artikelnaam '{description}' niet gevonden in de artikelentabel.")

    if data:
        new_df = pd.DataFrame(data, columns=[
            "Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", 
            "Aantal", "M2 p/s", "M2 totaal", "RSP", "SAP Prijs", "Handmatige Prijs", 
            "Min_prijs", "Max_prijs", "Verkoopprijs", "Prijs_backend", "Source", "fuzzy_match", "original_article_number"
        ])
        st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_df], ignore_index=True)
        st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, prijsscherpte)
        st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
        st.rerun()
    else:
        st.sidebar.warning("Geen gegevens gevonden om toe te voegen.")

def remap_and_process(df):
    # Hier kun je logica toevoegen om de achtergehouden regels opnieuw te mappen
    st.write("Her-mapping van achtergehouden regels...")
    return df



def manual_column_mapping(df, detected_columns):
    """
    Biedt de gebruiker een interface om ontbrekende kolommen handmatig te mappen,
    waarbij JSON-kolommen worden gebruikt als deze beschikbaar zijn.
    """

    # **Stap 1: Gebruik JSON DataFrame als die bestaat**
    if "json_df" in st.session_state and st.session_state["json_df"] is not None and not st.session_state["json_df"].empty:
        df = st.session_state["json_df"].copy()  # Gebruik JSON-data
        st.write("📌 **Data geladen vanuit JSON-extractie**")
    else:
        st.write("📌 **Geen JSON-data gevonden, gebruik standaard dataset**")

    all_columns = list(df.columns)  # **Gebruik alleen kolommen van de gekozen dataset**
    st.write("🛠 **Beschikbare kolommen:**", all_columns)  # Debugging stap

    mapped_columns = {}  # Lege mapping dictionary

    st.write("Controleer of de kolommen correct zijn gedetecteerd.✨ Indien niet, selecteer de juiste kolom.")

    for key in ["Artikelnaam", "Hoogte", "Breedte", "Aantal"]:
        session_key = f"mapped_{key}"

        # **Stap 2: Initialiseer session_state vóór het aanmaken van de widget**
        if session_key not in st.session_state:
            st.session_state[session_key] = "Geen"

        # **Stap 3: Correcte index bepalen**
        options = ["Geen"] + all_columns
        default_index = options.index(st.session_state[session_key]) if st.session_state[session_key] in options else 0

        # **Stap 4: Selectbox met JSON-kolommen als deze beschikbaar zijn**
        mapped_columns[key] = st.selectbox(
            f"Selecteer kolom voor '{key}'",
            options=options,
            index=default_index,
            key=session_key
        )

    # **Stap 5: Filter de mapping om alleen daadwerkelijke selecties te behouden**
    mapped_columns = {k: v for k, v in mapped_columns.items() if v != "Geen"}

    return mapped_columns




# Functie voor PDF naar Excel conversie
def pdf_to_excel(pdf_reader, excel_path):
    try:
        with pdfplumber.open(pdf_reader) as pdf:
            writer = pd.ExcelWriter(excel_path, engine='openpyxl')
            has_data = False

            for i, page in enumerate(pdf.pages):
                table = page.extract_table()
                if table and len(table) > 1:
                    headers = table[0] if all(isinstance(h, str) for h in table[0]) else [f"Kolom_{j}" for j in range(len(table[0]))]
                    df = pd.DataFrame(table[1:], columns=headers)
                    if not df.empty:
                        df.to_excel(writer, sheet_name=f"Page_{i+1}", index=False)
                        has_data = True

            if has_data:
                writer.close()
                return excel_path
            else:
                writer.close()
                return None

    except Exception as e:
        pass
        return None




def is_valid_numeric(value, min_value):
    """ Controleert of een waarde numeriek is en groter dan een minimale waarde. """
    try:
        num = float(value)
        return num > min_value
    except (ValueError, TypeError):
        return False

def shift_row_left(row_values, start_index, shift_amount):
    """ Schuift alle waarden rechts van start_index naar links met shift_amount. """
    new_row = row_values.copy()
    new_row[start_index:-shift_amount] = new_row[start_index+shift_amount:]
    new_row[-shift_amount:] = None  # Maak de laatste kolommen leeg na verschuiving
    return new_row

def correct_backlog_rows(df_backlog):
    """
    Corrigeer rijen die in de backlog zitten door de kolommen systematisch naar links te verschuiven
    vanaf de eerste None-waarde.
    """
    corrected_rows = []
    
    for _, row in df_backlog.iterrows():
        row_values = row.values.copy()
        none_index = np.where(pd.isna(row_values))[0]
        
        if len(none_index) > 0:
            none_col = none_index[0]  # Eerste None waarde gevonden
            
            for shift in [1, 2]:  # Probeer 1 en 2 kolommen naar links te schuiven
                corrected_row = shift_row_left(row_values, none_col, shift)
                corrected_series = pd.Series(corrected_row, index=df_backlog.columns)
                
                if (
                    is_valid_numeric(corrected_series["aantal"], 0) and
                    is_valid_numeric(corrected_series["breedte"], 99) and
                    is_valid_numeric(corrected_series["hoogte"], 99)
                ):
                    corrected_rows.append(corrected_series)
                    break
            else:
                corrected_rows.append(row)
        else:
            corrected_rows.append(row)
    
    return pd.DataFrame(corrected_rows, columns=df_backlog.columns)

def extract_text_from_pdf(pdf_bytes):
    """
    Haalt tekst uit een PDF-bestand.
    """
    try:
        with pdfplumber.open(pdf_bytes) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        return text
    except Exception as e:
        st.error(f"Fout bij tekstextractie uit PDF: {e}")
        return ""

def extract_pdf_to_dataframe(pdf_reader):
    try:
        # **Stap 1: Controleer of er een tabel in de PDF staat**
        table_found = False  # Flag om bij te houden of een tabel is gevonden
        first_table = None  # Variabele om eerste gevonden tabel op te slaan

        with pdfplumber.open(pdf_reader) as pdf:
            for i, page in enumerate(pdf.pages):
                table = page.extract_table()
                if table:
                    table_found = True  # Markeer dat er een tabel is gevonden
                    first_table = table  # Sla de eerste gevonden tabel op
                    break  # Eén tabel is genoeg om de check te voltooien
        
        # **Toon de uitkomst in de UI**
        if table_found and first_table:
            st.success("✅ Een tabel is gevonden in de PDF.")
            df_table = pd.DataFrame(first_table[1:], columns=first_table[0])  # Eerste rij als header gebruiken
            
            # **Debugging Stap**: Controleer of er duplicate indexwaarden zijn
            st.write("📌 **Debugging: Inhoud van df_table vóór index reset**")
            st.write(df_table)

            if df_table.index.duplicated().any():
                st.error("⚠ Waarschuwing: Dubbele indexen gedetecteerd in de tabel!")
                df_table = df_table.reset_index(drop=True)  # Fix index probleem
            
            st.write("📌 **Debugging: DataFrame na index reset**")
            st.write(df_table)

            st.write("**Voorbeeld van de eerste gedetecteerde tabel:**")
            st.dataframe(df_table)  # Toon de tabel in de UI
            return df_table  # Return de tabel als dataframe
        else:
            if not table_found:
                st.warning("✨ Geen tabel gevonden.")
            
                        
                # Voer nu de AI-extractie uit
                document_text = extract_text_from_pdf(pdf_reader)
                relevant_data = extract_data_with_gpt(document_text)

                # **Stap 1: Bewaar JSON-output in session_state**
                if "json_df" not in st.session_state or st.session_state["json_df"] is None:
                    st.session_state["json_df"] = relevant_data.copy()  # Sla AI-extractie op
                
                # Verwijder de progress bar en geef succesmelding
                progress_bar.empty()
                st.success("✅ AI-extractie voltooid!")
                # **Debugging: Toon ruwe GPT-response**
                st.write("📌 **Debugging: Ruwe GPT-response (exacte output van GPT)**")
                st.code(relevant_data, language="json")
                
                # **Controleer of de respons een geldige DataFrame is**
                if isinstance(relevant_data, pd.DataFrame) and not relevant_data.empty:
                    st.success("✅ AI-extractie voltooid!")
                    st.write("📌 **Data geëxtraheerd via AI:**")
                    st.dataframe(relevant_data)
                    return relevant_data  # Direct GPT-resultaat retourneren
                else:
                    st.error("❌ Fout bij GPT-extractie: De gegenereerde data is niet geldig.")
                    return pd.DataFrame()  # Voorkom crashes door een lege DataFrame terug te geven
                
                st.success("✅ AI-extractie voltooid!")
                st.write("📌 **Data geëxtraheerd via AI:**")
                st.dataframe(relevant_data)
                return relevant_data  # Direct GPT-resultaat retourneren

            else:
                st.warning("⚠ Geen tabel gevonden, en AI-extractie is niet ingeschakeld.")


            # **Fallback naar tekstextractie als er geen tabel is gevonden**
            with pdfplumber.open(pdf_reader) as pdf:
                lines = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        lines.extend(text.split("\n"))
    
            structured_data = []
            current_category = "0"  # Fallback waarde als er geen categorie is
            category_pattern = re.compile(r"(\d{1,2}-\s*\d{1,2}[A-Z]?-\s*\w+)|(\d{1,2}[A-Z]?\s*-\s*\w+)")

            for line in lines:
                line = line.strip()
                if category_pattern.match(line):
                    current_category = line.replace(":", "")
                    continue
    
                if re.search(r"\bTotaal:?\b", line, re.IGNORECASE):
                    continue
    
                columns = re.split(r'\s+', line)
                if len(columns) >= 5:
                    structured_data.append([current_category] + columns)
    
            if structured_data:
                max_columns = max(len(row) for row in structured_data)
                column_names = ["Categorie"] + [f"Kolom_{i}" for i in range(1, max_columns)]
                structured_data = [row + [""] * (max_columns - len(row)) for row in structured_data]
                df = pd.DataFrame(structured_data, columns=column_names)
    
                header_row = None
                for i in range(min(3, len(df))):  # Maximaal 3 rijen doorzoeken voor headers
                    potential_headers = df.iloc[i].astype(str).str.lower().str.strip()
                    if any(potential_headers.isin([
                        "artikelnaam", "artikel", "product", "type", "article", "samenstelling",
                        "hoogte", "height", "h",
                        "breedte", "width", "b",
                        "aantal", "quantity", "qty", "stuks"
                    ])):
                        header_row = i
                        break
    
                if header_row is not None:
                    # Controleer en hernoem dubbele kolomnamen
                    df.columns = df.iloc[header_row].astype(str).str.strip()  # Strip spaties en zet om naar string
                    df = df.drop(df.index[:header_row + 1])  # Verwijder header-rij
                    
                    # Hernoem dubbele kolommen
                    if df.columns.duplicated().any():
                        st.error(f"⚠ Fout: Dubbele kolomnamen gevonden: {df.columns[df.columns.duplicated()].tolist()}")
                    
                        new_columns = []
                        col_count = {}
                        for col in df.columns:
                            if col in col_count:
                                col_count[col] += 1
                                new_columns.append(f"{col}_{col_count[col]}")  # Voeg index toe aan dubbele kolommen
                            else:
                                col_count[col] = 1
                                new_columns.append(col)
                    
                        df.columns = new_columns  # Update kolomnamen
                        st.success("✅ Dubbele kolomnamen hernoemd.")

                    df = df.drop(df.index[:header_row + 1])
                else:
                    st.warning("⚠ Geen header herkend, eerste rij als header gebruikt.")
                    df.columns = df.iloc[0]
                    df = df.drop(df.index[0])
                
                # **Debugging Stap**: Controleer of de index uniek is
                st.write("📌 **Debugging: Inhoud van df vóór index reset**")
                st.write(df)

                if not df.index.is_unique:
                    st.error("⚠ Waarschuwing: Niet-unieke indexwaarden gevonden vóór reset. Fix index...")
                    st.write("Dubbele indexen gevonden:", df.index[df.index.duplicated()].tolist())  # Debug info
                    df = df.loc[~df.index.duplicated(keep='first')].reset_index(drop=True)
                else:
                    df = df.reset_index(drop=True)

                # Debugging: Controleer opnieuw na reset
                if not df.index.is_unique:
                    st.error("⚠ Probleem na reset: Index is nog steeds niet uniek!")
                    st.write("Huidige indexstatus:", df.index)

                # Extra: Print de kolommen en rijen om te checken of data correct is
                st.write("📌 **Debugging: DataFrame na index reset**")
                st.write(df)

                df.columns = df.columns.str.lower()
    
                if "aantal" not in df.columns:
                    st.error("⚠ Kolom 'aantal' niet gevonden in de PDF.")
                    st.write("Herkende kolommen:", df.columns.tolist())
                    return pd.DataFrame()
    
                if df.shape[0] > 2:
                    df = pd.concat([
                        df.iloc[:2],  
                        df.iloc[2:][~df.iloc[2:].apply(
                            lambda row: row.astype(str).str.contains(r"\b(Aantal|Breedte|Hoogte)\b", case=False).any(), axis=1)
                        ]
                    ]).reset_index(drop=True)
    
                for col in ["aantal", "breedte", "hoogte"]:
                    if col in df.columns:
                        df[col] = pd.to_numeric(df[col], errors="coerce")  
                
                if "df_current" not in st.session_state:
                    st.session_state.df_current = df.copy()
                if "batch_number" not in st.session_state:
                    st.session_state.batch_number = 1
                if "next_df" not in st.session_state:
                    st.session_state.next_df = None
                
                if st.session_state.next_df is not None:
                    st.session_state.df_current = st.session_state.next_df.copy()
                    st.session_state.next_df = None  
                
                df_current = st.session_state.df_current
                
                df_backlog = df_current[
                    df_current["aantal"].isna() | (df_current["aantal"] <= 0) |
                    df_current["breedte"].isna() | (df_current["breedte"] < 100) |
                    df_current["hoogte"].isna() | (df_current["hoogte"] < 100)
                ]
                
                if not df_backlog.empty:
                    df_corrected = correct_backlog_rows(df_current)
                    df_current.update(df_corrected)
                
                df_bulk = df_current.loc[
                    ~df_current.index.isin(df_backlog.index)
                ].copy()
    
                st.write("✅ **Verwerkte gegevens:**")
                st.dataframe(df_corrected)
                
                return df_corrected  
    
            else:
                st.warning("Geen gegevens gevonden in de PDF om te verwerken test.")
    
    except Exception as e:
        st.error(f"Fout bij het extraheren van PDF-gegevens: {e}")
        return pd.DataFrame()






        
def extract_latest_email(body):
    """
    Extracts only the latest email from an email thread.
    It detects the start of a new email using the pattern 'Van:' followed by 'Verzonden:'.
    """
    email_parts = re.split(r'Van:.*?Verzonden:.*?Aan:.*?Onderwerp:', body, flags=re.DOTALL)
    if email_parts:
        latest_email = email_parts[0].strip()
        return latest_email
    else:
        return body.strip()



def debug_check_tables(doc_bytes):
    """ Controleert of er tabellen in het DOCX-bestand zijn en toont extra statistieken. """
    doc = Document(BytesIO(doc_bytes))
    num_tables = len(doc.tables)

    print(f"📊 Aantal tabellen in het DOCX-bestand: {num_tables}")

    total_table_cells = 0
    total_text_lines = 0
    total_words = 0
    total_chars = 0

    # **Tabelinformatie**
    for i, table in enumerate(doc.tables):
        st.write(f"📂 Tabel {i+1}:")
        for row in table.rows:
            cell_values = [cell.text.strip() for cell in row.cells]
            st.write(cell_values)
            total_table_cells += len(cell_values)
        st.write("="*50)  # Visuele scheiding tussen tabellen

    if num_tables == 0:
        st.write("❌ Geen tabellen gevonden in het document!")

    # **Tekstinformatie**
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            total_text_lines += 1
            total_chars += len(text)
            total_words += len(text.split())

    # **Toon de statistieken**
    st.write(f"📃 Aantal regels tekst: {total_text_lines}")
    st.write(f"🔤 Aantal tekens: {total_chars}")
    st.write(f"📝 Aantal woorden: {total_words}")

    if num_tables > 0:
        st.write(f"📦 Totaal aantal cellen in tabellen: {total_table_cells}")

def convert_docx_to_xlsx(doc_bytes):
    """
    Converteer een DOCX-bestand naar een Excel-bestand en neem ALLE inhoud mee.
    """
    # **Voer eerst de debug-check uit**
    debug_check_tables(doc_bytes)

    # Maak een tijdelijk Excel-bestand aan
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
        excel_output_path = temp_file.name

    with pd.ExcelWriter(excel_output_path, engine="openpyxl") as writer:
        table_count = 0
        has_data = False

        # **Stap 1: Controleer tabellen**
        if len(doc.tables) > 0:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):  # Lege rijen negeren
                        rows.append(cells)

                if rows:
                    df = pd.DataFrame(rows)
                    df.columns = [f"Kolom_{i+1}" for i in range(df.shape[1])]  # Fallback headers

                    table_count += 1
                    df.to_excel(writer, sheet_name=f"Tabel_{table_count}", index=False)
                    has_data = True
        else:
            st.write("❌ Geen tabellen gevonden! We proberen tekst als tabel te verwerken.")

        # **Stap 2: Als er geen tabellen zijn, probeer tekstregels als tabel te extraheren**
        structured_data = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Probeer regels te splitsen op tab of meerdere spaties (lijkt op een tabel)
                row = [t.strip() for t in text.split("\t")]  # Eerst op tabs splitsen
                if len(row) < 2:  # Als te weinig kolommen, probeer spaties
                    row = [t.strip() for t in text.split("  ")]  # Twee of meer spaties als scheiding
                structured_data.append(row)

        # **Stap 3: Sla gestructureerde tekst op als tabel**
        if structured_data:
            df_text = pd.DataFrame(structured_data)
            df_text.columns = [f"Kolom_{i+1}" for i in range(df_text.shape[1])]  # Fallback headers
            df_text.to_excel(writer, sheet_name="Gestructureerde Tekst", index=False)
            has_data = True

        # **Stap 4: Voeg een minimale zichtbare sheet toe als er geen gegevens zijn**
        if not has_data:
            df_empty = pd.DataFrame({"Melding": ["Geen data gevonden"]})
            df_empty.to_excel(writer, sheet_name="Leeg Document", index=False)

    return excel_output_path


def extract_data_with_gpt(prompt):
    """
    Verstuurt een tekstprompt naar GPT en retourneert een correct geformatteerde DataFrame.
    """
    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": (
                    "Je bent een geavanceerde extractietool die glassamenstellingen uit een bestekformulier extraheert "
                    "en deze omzet naar een correcte JSON-tabel. "
                    "Zorg ervoor dat de JSON-structuur voldoet aan de volgende vereisten:\n\n"
                    "1️ **Elke regel in de JSON moet minstens een 'glasType' of 'omschrijving' van het artikel, de 'hoogte', de 'breedte' en het 'aantal' bevatten** Vind je geen glastype om omschrijving van het artikel? Pak dan het artikel van de voorgaande regel.\n"
                    "2️ **'aantal', 'breedte' en 'hoogte' moeten op het hoofdniveau staan** en mogen NIET in de 'details'-array of geneste functie geplaatst worden.\n"
                    "3 **De JSON-output mag GEEN extra uitleg bevatten**, enkel en alleen de gestructureerde JSON-data.\n\n"
                    "Geef de output als een **geldige JSON-array**, zonder extra tekst, uitleg of Markdown-codeblokken."
                )},
                {"role": "user", "content": prompt}
            ]
        )

        
        extracted_data = response.choices[0].message.content  # Haal de tekstuele GPT-output op

        # **Stap 1: Debugging - Toon ruwe GPT-response**
        st.write("📌 **Debugging: Ruwe GPT-response (exacte output van GPT)**")
        st.code(extracted_data, language="json")

        # **Stap 2: Strip Markdown en onnodige tekens**
        extracted_data = extracted_data.strip()
        if extracted_data.startswith("```json"):
            extracted_data = extracted_data[7:].strip()  # Verwijder '```json'
        if extracted_data.endswith("```"):
            extracted_data = extracted_data[:-3].strip()  # Verwijder '```'

        # **Stap 3: JSON validatie**
        try:
            extracted_json = json.loads(extracted_data)  # Parse JSON
        except json.JSONDecodeError:
            st.error("❌ GPT-response is geen geldige JSON! Controleer de AI-output.")
            return pd.DataFrame()  # Lege DataFrame als fallback

        # **Stap 4: Zet JSON om naar een DataFrame**
        if isinstance(extracted_json, list):
            df_json = pd.DataFrame(extracted_json)
        elif isinstance(extracted_json, dict):
            df_json = pd.DataFrame([extracted_json])  # Zet een enkele dict om naar DataFrame
        else:
            st.error("❌ GPT-response heeft geen correct formaat.")
            return pd.DataFrame()  # Leeg DataFrame als fallback

        # **Stap 5: Converteer numerieke kolommen**
        for col in df_json.columns:
            if df_json[col].dtype == "object":  # Alleen stringkolommen aanpassen
                df_json[col] = df[col].astype(str).str.replace(" mm", "", regex=True)
                df_json[col] = df[col].astype(str).str.replace(" m²", "", regex=True)

                # Probeer te converteren naar numeriek indien mogelijk
                df_json[col] = pd.to_numeric(df[col], errors="ignore")

        # **Stap 6: Toon de verwerkte DataFrame**
        st.success("✅ AI-extractie voltooid! Hieronder de geformatteerde output:")
        st.dataframe(df_json)
        return df_json

    except Exception as e:
        st.error(f"❌ Fout bij GPT-extractie: {e}")
        return pd.DataFrame()  # Leeg DataFrame als fallback


    

def process_attachment(attachment, attachment_name):
    """
    Verwerkt een bijlage op basis van het bestandstype (Excel of PDF) en past automatisch kolommapping toe.
    """

    # Bestandstypes die GEEN knop mogen krijgen
    excluded_extensions = ('.png', '.jpg', '.jpeg')

    # **Toon GEEN knop voor afbeeldingen**
    if attachment_name.lower().endswith(excluded_extensions):
        return  # Stop de functie voor deze bestandstypes

    # **AI Extractie alleen als gebruiker op de knop drukt**
    if st.sidebar.button(f"🔍 Gebruik AI-extractie voor {attachment_name}"):
        with st.spinner(f"AI-extractie bezig voor {attachment_name}... ⏳"):
            if attachment_name.endswith(".pdf"):
                document_text = extract_text_from_pdf(attachment)
            elif attachment_name.endswith(".xlsx"):
                document_text = extract_text_from_excel(attachment)
            else:
                st.error("❌ Dit bestandsformaat wordt niet ondersteund voor AI-extractie.")
                return

            relevant_data = extract_data_with_gpt(document_text)

            # **Sla JSON-output op in session_state**
            if isinstance(relevant_data, pd.DataFrame) and not relevant_data.empty:
                st.session_state["json_df"] = relevant_data.copy()
                st.success("✅ AI-extractie voltooid!")

    # **Gebruik JSON als deze al bestaat**
    if "json_df" in st.session_state and not st.session_state["json_df"].empty:
        df_extracted = st.session_state["json_df"]
        st.write("📌 **Data geladen vanuit AI-extractie**")
    else:
        # **Laad de data alleen als er nog geen AI-extractie is uitgevoerd**
        if attachment_name.endswith(".pdf"):
            df_extracted = extract_pdf_to_dataframe(attachment, True)  # AI-extractie niet meer nodig, al gebeurd
        elif attachment_name.endswith(".xlsx"):
            df_extracted = pd.read_excel(BytesIO(attachment), dtype=str)
        else:
            st.error("❌ Dit bestandsformaat wordt niet ondersteund.")
            return

    if not df_extracted.empty:
        detected_columns = detect_relevant_columns(df_extracted)
        mapped_columns = manual_column_mapping(df_extracted, detected_columns)

    if attachment_name.endswith(".xlsx"):
        try:
            df = pd.read_excel(BytesIO(attachment), dtype=str)  # Inlezen als strings
            st.write("Bijlage ingelezen als DataFrame:")
            st.dataframe(df)

            # Automatische header-detectie
            header_row = None
            for i in range(min(30, len(df))):  # Zoek de header in de eerste 30 rijen
                potential_headers = df.iloc[i].fillna("").astype(str).str.lower().str.strip()  # Voorkom 'float' errors
                if any(potential_headers.isin([
                    "artikelnaam", "artikel", "product", "type", "article", "samenstelling",
                    "hoogte", "height", "h",
                    "breedte", "width", "b",
                    "aantal", "quantity", "qty", "stuks"
                ])):
                    header_row = i
                    break

            if header_row is not None:
                df.columns = df.iloc[header_row].fillna("").astype(str).str.lower().str.strip()  # Kolomnamen corrigeren
                df = df.drop(df.index[:header_row + 1]).reset_index(drop=True)
            else:
                st.warning("Geen headers gedetecteerd in de eerste 30 rijen.")
                return None

            df.columns = df.columns.str.lower()

            # Verwijder onnodige rijen zoals 'Totaal'
            df = df[~df.apply(lambda row: row.fillna("").astype(str).str.contains(r'totaal', case=False).any(), axis=1)]
            df = df.dropna(how='all')

            # Detecteer en map kolommen
            detected_columns = detect_relevant_columns(df)
            mapped_columns = manual_column_mapping(df, detected_columns)

            if not isinstance(mapped_columns, dict):
                st.error("Mapping fout: mapped_columns is geen dictionary. Controleer de kolommapping.")
                return None

            if mapped_columns:
                relevant_data = df[[mapped_columns[key] for key in mapped_columns]]
                relevant_data.columns = mapped_columns.keys()

                # Filter relevant data
                start_row = st.sidebar.number_input("Beginrij data (niet de header):", min_value=0, max_value=len(df)-1, value=0)
                end_row = st.sidebar.number_input("Eindrij data:", min_value=0, max_value=len(df)-1, value=len(df)-1)
                relevant_data = relevant_data.iloc[int(start_row):int(end_row)+1]

                # GPT Extractie als geen data is gevonden
                if relevant_data.empty:
                    st.warning("Geen tabel gevonden. Probeer GPT-extractie...")

                    document_text = extract_text_from_excel(attachment)
                    if document_text:
                        relevant_data = extract_data_with_gpt(document_text)
                        st.write("Data geëxtraheerd via GPT:")
                        st.dataframe(relevant_data)
                    
                    if not relevant_data.empty and st.sidebar.button("Verwerk gegevens naar offerte"):
                        handle_mapped_data_to_offer(relevant_data)

                st.write("Relevante data:")
                st.dataframe(relevant_data)

                if st.sidebar.button("Verwerk gegevens naar offerte"):
                    handle_mapped_data_to_offer(relevant_data)

            else:
                st.warning("Geen relevante kolommen gevonden of gemapped.")
                return None
        except Exception as e:
            st.error(f"Fout bij het verwerken van de Excel-bijlage: {e}")
            return None


    elif attachment_name.endswith(".pdf"):
        try:
            pdf_reader = BytesIO(attachment)
            excel_path = "converted_file.xlsx"
    
            # PDF omzetten naar Excel
            excel_result = pdf_to_excel(pdf_reader, excel_path)
            if not excel_result:
                pass
    
            # Gegevens extraheren uit PDF
            df_extracted = extract_pdf_to_dataframe(pdf_reader)
            if not df_extracted.empty:

                # Verwijder onnodige rijen (zoals 'Totaal'-rijen)
                df_extracted = df_extracted[~df_extracted.apply(lambda row: row.astype(str).str.contains(r'totaal', case=False).any(), axis=1)]
                df_extracted = df_extracted.dropna(how='all')
                
                # Relevante kolommen detecteren
                detected_columns = detect_relevant_columns(df_extracted)
                mapped_columns = manual_column_mapping(df_extracted, detected_columns)
    
                if not mapped_columns:
                    pass
                    return
    
                # Selecteer en hernoem kolommen op basis van mapping
                relevant_data = df_extracted[list(mapped_columns.values())]
                relevant_data.columns = list(mapped_columns.keys())

                # GPT Extractie als geen data is gevonden
                if relevant_data.empty:
                    st.warning("Geen tabel gevonden. Probeer GPT-extractie...")

                    document_text = extract_text_from_pdf(attachment)
                    if document_text:
                        relevant_data = extract_data_with_gpt(document_text)
                        st.write("Data geëxtraheerd via GPT:")
                        st.dataframe(relevant_data)
                    
                    if not relevant_data.empty and st.sidebar.button("Verwerk gegevens naar offerte"):
                        handle_mapped_data_to_offer(relevant_data)
                
                st.write("Data na mapping:")
                st.dataframe(relevant_data)
    
                # Optionele filtering op rijen
                start_row = st.sidebar.number_input("Beginrij (inclusief):", min_value=0, max_value=len(relevant_data)-1, value=0)
                end_row = st.sidebar.number_input("Eindrij (inclusief):", min_value=0, max_value=len(relevant_data)-1, value=len(relevant_data)-1)
    
                relevant_data = relevant_data.iloc[int(start_row):int(end_row)+1]
    
                # Verwerken van de gegevens
                if not relevant_data.empty:
                    if st.button("Verwerk gegevens naar offerte"):
                        handle_mapped_data_to_offer(relevant_data)
                else:
                    st.warning("Relevante data is leeg. Controleer de kolommapping en inhoud van de PDF.")
            else:
                st.warning("Geen gegevens gevonden in de PDF om te verwerken. test2")
    
        except Exception as e:
            st.error(f"Fout bij het verwerken van de PDF-bijlage: {e}")


    elif attachment_name.endswith(".docx"):
        try:
            st.write(f"DOCX-bestand '{attachment_name}' ingelezen.")
    
            # Converteer DOCX naar Excel
            excel_file = convert_docx_to_xlsx(attachment)
    
            # Lees de Excel-data in
            df_dict = pd.read_excel(excel_file, sheet_name=None)
    
            # Controleer of er data is
            if not df_dict:
                st.error("Geen data gevonden in het Excel-bestand.")
                return None
    
            # Toon de beschikbare tabellen in de Excel
            st.write("Selecteer de tabel om te verwerken:")
            selected_sheet = st.sidebar.selectbox(
                "Kies een tabel:", options=list(df_dict.keys()), format_func=lambda x: f"Tabel: {x}"
            )
            table = df_dict[selected_sheet]
    
            st.write(f"Inhoud van **{selected_sheet}**:")
            st.dataframe(table)
    
            # **Detecteer relevante kolommen**
            detected_columns = detect_relevant_columns(table)
            mapped_columns = manual_column_mapping(table, detected_columns)
    
            # **Validatie voor mapped_columns**
            if not isinstance(mapped_columns, dict):
                st.error("Mapping fout: mapped_columns is geen dictionary. Controleer de kolommapping.")
                return None
    
            if mapped_columns:
                # **Selecteer en hernoem relevante kolommen**
                relevant_data = table[[mapped_columns[key] for key in mapped_columns]]
                relevant_data.columns = mapped_columns.keys()
    
                # **Filter relevante data op rijen**
                start_row = st.sidebar.number_input("Beginrij (inclusief):", min_value=0, max_value=len(table)-1, value=0)
                end_row = st.sidebar.number_input("Eindrij (inclusief):", min_value=0, max_value=len(table)-1, value=len(table)-1)
                relevant_data = relevant_data.iloc[int(start_row):int(end_row)+1]

                # **GPT Extractie als geen data is gevonden**
                if relevant_data.empty:
                    st.warning("Geen tabel gevonden. Probeer GPT-extractie...")

                    document_text = extract_text_from_docx(attachment)
                    if document_text:
                        relevant_data = extract_data_with_gpt(document_text)
                        st.write("Data geëxtraheerd via GPT:")
                        st.dataframe(relevant_data)

                    if not relevant_data.empty and st.sidebar.button("Verwerk gegevens naar offerte"):
                        handle_mapped_data_to_offer(relevant_data)
                
                
                st.write("Relevante data:")
                st.dataframe(relevant_data)
    
                if not relevant_data.empty:
                    if st.sidebar.button("Verwerk gegevens naar offerte"):
                        handle_mapped_data_to_offer(relevant_data)
                else:
                    st.warning("Relevante data is leeg. Controleer de kolommapping en inhoud van de tabel.")
            else:
                st.warning("Geen relevante kolommen gevonden of gemapped.")
    
        except Exception as e:
            st.error(f"Fout bij het verwerken van de DOCX-bijlage: {e}")


st.sidebar.markdown("---")  # Scheidingslijn voor duidelijkheid  

# File uploader alleen beschikbaar in de uitklapbare invoeropties
with st.sidebar.expander("Upload document", expanded=False):
    # Bestand uploaden
    uploaded_file = st.file_uploader("Upload een Outlook, PDF of Excel bestand", type=["msg", "pdf", "xlsx", "docx"])
    
    # Controleren of er een bestand is geüpload
    if uploaded_file:
        # Bestand tijdelijk opslaan
        with open("uploaded_email.msg", "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Open het .msg-bestand met extract-msg
        try:
            msg = extract_msg.Message("uploaded_email.msg")
            msg_subject = msg.subject
            msg_sender = msg.sender
            full_email_body = msg.body  # De volledige e-mailthread
            latest_email = extract_latest_email(full_email_body)  # Bepaal alleen de laatste e-mail
            msg_body = latest_email
            email_body = msg_body

            # Stel onderwerp van de mail in als klantreferentie als deze nog leeg is. Verwijder FW: of RE: vooraan het onderwerp.
            if msg_subject:
                try:
                    if not st.session_state.get("customer_reference") or not customer_reference.strip():
                        # Verwijder "FW: " of "RE: " aan het begin van msg_subject
                        clean_subject = re.sub(r"^(FW:|RE:)\s*", "", msg_subject.strip(), flags=re.IGNORECASE)
                        
                        # Gebruik het opgeschoonde onderwerp als klantreferentie
                        st.session_state["customer_reference"] = clean_subject
                        
                        # Trigger herladen van de interface
                        st.rerun()
                except Exception as e:
                    st.error(f"Fout bij het verwerken van de klantreferentie: {e}")


            
            # Resultaten weergeven
            st.subheader("Berichtinformatie")
            st.write(f"**Onderwerp:** {msg_subject}")
            st.write(f"**Afzender:** {msg_sender}")
            st.write("**Inhoud van het bericht:**")
            st.text(msg_body)
            
            # Verwerk bijlagen
            st.subheader("Bijlagen:")
            if msg.attachments:
                for attachment in msg.attachments:
                    attachment_name = attachment.longFilename or attachment.shortFilename
                    attachment_data = attachment.data
    
                    # Toon naam van de bijlage
                    st.write(f"Bijlage: {attachment_name}")
    
                    # Verwerk de bijlage
                    process_attachment(attachment_data, attachment_name)
            else:
                st.write("Geen bijlagen gevonden.")
        
        except Exception as e:
            st.error(f"Fout bij het verwerken van het bestand: {e}")
    else:
        st.info("Upload een .msg-bestand om verder te gaan.") 


# Gebruikersinvoer
customer_input = st.sidebar.text_area("Voer hier handmatig het klantverzoek in.")


# Functie om tekstinvoer te verwerken
def handle_text_input(input_text):
    matched_articles = [(term, synonym_dict[term]) for term in synonym_dict if term in input_text]

    if matched_articles:
        response_text = "Bedoelt u de volgende samenstellingen:"
        for term, article_number in matched_articles:
            description, _, _, _, _ = find_article_details(article_number)
            if description:
                response_text += f"- {description} met artikelnummer {article_number}\n"

        response_text += "?"
        st.sidebar.write(response_text)
    else:
        st.sidebar.warning("Geen gerelateerde artikelen gevonden. Gelieve meer details te geven.")

# Functie om offerte als PDF te genereren
def generate_pdf(df):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from io import BytesIO

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []

    # Styles
    styles = getSampleStyleSheet()
    header_style = ParagraphStyle(
        'HeaderStyle', parent=styles['Heading1'], fontSize=17, alignment=TA_CENTER, textColor=colors.black
    )
    normal_style = ParagraphStyle(
        'NormalStyle', parent=styles['Normal'], fontSize=11, alignment=TA_LEFT, textColor=colors.black
    )
    right_aligned_style = ParagraphStyle(
        'RightAlignedStyle', parent=styles['Normal'], fontSize=11, alignment=TA_LEFT, textColor=colors.black
    )

    # Header
    elements.append(Paragraph("Vandaglas - Offerte", header_style))
    elements.append(Spacer(1, 12))

    # Introductietekst
    elements.append(Paragraph(
        "Beste klant,<br/><br/>"
        "Hartelijk dank voor uw prijsaanvraag. Hieronder vindt u onze offerte. Wij hopen u een passend aanbod te hebben gedaan. "
        "Uw contactpersoon, Job, geeft graag nog een toelichting en beantwoordt eventuele vragen.<br/><br/>"
        "Met vriendelijke groet,<br/>"
        "Vandaglas",
        normal_style
    ))
    elements.append(Spacer(1, 24))

    # Tabel header
    data = [["Artikelnaam", "Breedte", "Hoogte", "Aantal", "Prijs p/s", "M2 p/s", "Totaal M2", "Totaal"]]

    # Voeg gegevens uit df toe aan tabel
    for index, row in df.iterrows():
        if all(col in row for col in ['Artikelnaam', 'Breedte', 'Hoogte', 'Aantal', 'RSP', 'M2 p/s', 'M2 totaal']):
            data.append([
    row['Artikelnaam'],
    row['Breedte'],
    row['Hoogte'],
    row['Aantal'],
    row['Prijs_backend'],
    f"{float(str(row['M2 p/s']).replace('m²', '').replace(',', '.').strip()):.2f} m2" if pd.notna(row['M2 p/s']) else None,
    f"{float(str(row['M2 totaal']).replace('m²', '').replace(',', '.').strip()):.2f} m2" if pd.notna(row['M2 totaal']) else None,
    f"{round(float(str(row['Prijs_backend']).replace('€', '').replace(',', '.').strip()) * float(row['Aantal']) * float(str(row['M2 p/s']).replace('m²', '').replace(',', '.').strip()), 2):,.2f}" if pd.notna(row['Prijs_backend']) and pd.notna(row['Aantal']) else None
])


    # Maak de tabel
    table = Table(data, repeatRows=1, colWidths=[150, 45, 45, 45, 45, 45, 45, 60])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.black),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('WORDWRAP', (0, 0), (-1, -1), True),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 24))

    # Eindtotaal, BTW, Te betalen
    total_price = df.apply(lambda row: round(float(str(row['Prijs_backend']).replace('€', '').replace(',', '.').strip()) * float(str(row['M2 totaal']).replace('m²', '').replace(',', '.').strip()), 2) if pd.notna(row['Prijs_backend']) and pd.notna(row['M2 totaal']) else 0, axis=1).sum()
    btw = total_price * 0.21
    te_betalen = total_price + btw

    # Maak klein tabelletje voor totalen
    totals_data = [
        ["Eindtotaal:", f"€ {total_price:.2f}"],
        ["BTW (21%):", f"€ {btw:.2f}"],
        ["Te betalen:", f"€ {te_betalen:.2f}"]
    ]
    totals_table = Table(totals_data, colWidths=[100, 100])
    totals_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.black),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))

    from reportlab.lib.units import inch
    elements.append(Spacer(1, 0.5 * inch))
    totals_table = Table(totals_data, colWidths=[100, 100], hAlign='RIGHT')
    totals_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.black),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    elements.append(Spacer(1, 3 * inch))
    elements.append(totals_table)
  

    # Bouwelementen aan document
    doc.build(elements)
    buffer.seek(0)
    return buffer

# # Offerte Genereren tab
# with tab1:
#     # Knop om GPT-chat te versturen
#     if st.sidebar.button("Vertaal chat naar offerte"):
#         try:
#             handle_gpt_chat()
#         except Exception as e:
#             st.sidebar.error(f"Er is een fout opgetreden: {e}")

#     # Knop om de e-mail te vertalen naar een offerte
#     if st.sidebar.button("Vertaal mail naar offerte"):
#         try:
#             handle_email_to_offer(email_body)
#         except Exception as e:
#             st.error(f"Fout bij het verwerken van de e-mail: {e}")

# Offerte Genereren tab
with tab1:
    # Een container voor de knop en de informatiebalk
    with st.sidebar.container():   
        relevant_data = None

        # Verwerk de bijlage zodra deze is geüpload
        if uploaded_file is not None:
            attachment_name = uploaded_file.name
            relevant_data = process_attachment(uploaded_file.getvalue(), attachment_name)
        
        # Eén knop om de acties uit te voeren
        if st.sidebar.button("BullsAI 🚀"):
            actie_uitgevoerd = False
        
            # Spinner toevoegen rond alle acties
            with st.spinner("BullsAI is bezig met de verwerking..."):
                # Probeer de eerste actie (tekstvak naar offerte)
                try:
                    handle_gpt_chat()
                    actie_uitgevoerd = True
                except Exception:
                    pass  # Fout negeren en doorgaan naar de volgende actie
        
                # Als de eerste actie niet slaagt, probeer de tweede (bijlage mail)
                if not actie_uitgevoerd and relevant_data is not None:
                    try:
                        handle_mapped_data_to_offer(relevant_data)
                        actie_uitgevoerd = True
                    except Exception:
                        pass  # Fout negeren en doorgaan naar de volgende actie
        
                # Als de tweede actie niet slaagt, probeer de derde (mail naar offerte)
                if not actie_uitgevoerd:
                    try:
                        handle_email_to_offer(email_body)
                        actie_uitgevoerd = True
                    except Exception:
                        pass  # Fout negeren
        
            # Eindstatus bepalen
            if actie_uitgevoerd:
                pass
            else:
                st.error("BullsAI heeft geen gegevens kunnen verwerken.")

# Voeg rijnummers toe aan de offerte DataFrame als deze nog niet bestaat
if 'Rijnummer' not in st.session_state.offer_df.columns:
    st.session_state.offer_df.insert(0, 'Rijnummer', range(1, len(st.session_state.offer_df) + 1))

    


# Offerte Genereren tab
with tab1:    
    with col6:
        # Voeg een knop toe om de offerte als PDF te downloaden
        if totaal_bedrag > 25000:
            st.button("Download offerte als PDF", key='download_pdf_button', disabled=True)
            st.button("Autoriseer offerte", key='authorize_offer_button')
        else:
            if st.button("Download offerte als PDF", key='download_pdf_button'):
                pdf_buffer = generate_pdf(st.session_state.offer_df)
                st.download_button(label="Download PDF", data=pdf_buffer, file_name="offerte.pdf", mime="application/pdf")
    
        # Knop om offerte op te slaan in database
        if st.button("Sla offerte op"):
            try:
                # Haal de ingelogde Windows-gebruikersnaam op
                import os
                windows_user = getpass.getuser() if getpass.getuser() else "Onbekende gebruiker"
        
                # Zoek het hoogste offertenummer
                if not st.session_state.saved_offers.empty:
                    max_offer_number = st.session_state.saved_offers['Offertenummer'].max()
                    offer_number = max_offer_number + 1
                else:
                    offer_number = 1
        
                # Bereken eindtotaal
                if all(col in st.session_state.offer_df.columns for col in ['RSP', 'M2 totaal']):
                    eindtotaal = st.session_state.offer_df.apply(
                        lambda row: float(row['RSP']) * float(row['M2 totaal']) if pd.notna(row['RSP']) and pd.notna(row['M2 totaal']) else 0,
                        axis=1
                    ).sum()
                else:
                    eindtotaal = 0
        
                # Voeg offerte-informatie toe aan opgeslagen offertes in sessie
                offer_summary = pd.DataFrame({
                    'Offertenummer': [offer_number],
                    'Klantnummer': [str(st.session_state.customer_number)],
                    'Eindbedrag': [eindtotaal],
                    'Datum': [datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
                    'Gebruiker': [windows_user]  # Voeg gebruikersnaam toe
                })
                st.session_state.saved_offers = pd.concat([st.session_state.saved_offers, offer_summary], ignore_index=True)
        
                # Voeg offertenummer en gebruikersnaam toe aan elke regel in de offerte
                st.session_state.offer_df['Offertenummer'] = offer_number
                st.session_state.offer_df['Gebruiker'] = windows_user
        
                # Opslaan in database
                conn = create_connection()
        #        cursor = conn.cursor()
        
                try:
                    # Voeg elke rij van de offerte toe aan de database
                    for index, row in st.session_state.offer_df.iterrows():
                        cursor.execute("""
                        INSERT INTO Offertes (Offertenummer, Rijnummer, Artikelnaam, Artikelnummer, Spacer, Breedte, Hoogte, Aantal, 
                                              M2_per_stuk, M2_totaal, RSP, SAP_Prijs, Handmatige_Prijs, Min_prijs, Max_prijs, 
                                              Prijs_backend, Verkoopprijs, Source, Datum, Gebruiker)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """, (
                            row['Offertenummer'], index + 1, row['Artikelnaam'], row['Artikelnummer'], row['Spacer'],
                            row['Breedte'], row['Hoogte'], row['Aantal'], row['M2 p/s'], row['M2 totaal'], 
                            row['RSP'], row['SAP Prijs'], row['Handmatige Prijs'], row['Min_prijs'], row['Max_prijs'], 
                            row['Prijs_backend'], row['Verkoopprijs'], row['Source'], datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            windows_user  # Sla gebruikersnaam op
                        ))
        
                    conn.commit()
                    st.success(f"Offerte {offer_number} succesvol opgeslagen door {windows_user}.")
                except sqlite3.Error as e:
                    st.error(f"Fout bij het opslaan in de database: {e}")
                finally:
                    conn.close()
            except Exception as e:
                st.error(f"Er is een fout opgetreden: {e}")


if 'edited_df' in locals() and not edited_df.equals(st.session_state.offer_df):
    edited_df = edited_df.copy()
    edited_df = update_offer_data(edited_df)
    st.session_state.offer_df = edited_df


with tab1:
    # Checkbox voor het creëren van een Opportunity
    if st.checkbox("Creeer Opportunity"):
        # Indeling in vier kolommen
        col1, col2, col3, col4 = st.columns(4)
    
        # Velden in de eerste kolom
        with col1:
            name = st.text_input("Name (gevuld met customer_reference):", value=customer_reference)
            account_id = st.text_input("AccountID:", value="001KI0000084Q8VYAU")
            stage_name = st.selectbox(
                "StageName:",
                options=["RFQ / Initial Quote", "Customer is fixed", "Negotiation", "Verbal Agreement"],
                index=0,  # Standaard geselecteerde waarde
            )
            close_date = st.date_input(
                "CloseDate (datum vandaag + 2 weken):",
                value=date.today() + timedelta(weeks=2),
            )
            amount = st.number_input("Amount (gevuld met totaal_bedrag):", value=totaal_bedrag)
            description = st.text_area("Description (gevuld met customer_reference):", value=customer_reference)
    
        # Knop om de Opportunity aan te maken
        if st.button("Opportunity aanmaken"):
            try:
                # Opportunity-gegevens
                opportunity_data = {
                    "Name": name,
                    "AccountId": account_id,
                    "StageName": stage_name,
                    "CloseDate": close_date.isoformat(),
                    "Amount": amount,
                    "Description": description,
                }
    
                # Opportunity aanmaken in Salesforce
                resultaat = sf.Opportunity.create(opportunity_data)
    
                # Maak een hyperlink naar de Salesforce Opportunity
                opportunity_id = resultaat['id']
                salesforce_url = f"https://vandaglasnl--qa.sandbox.my.salesforce.com/{opportunity_id}"
                hyperlink = f"[{customer_reference}]({salesforce_url})"
    
                st.success(f"Opportunity succesvol aangemaakt! {hyperlink}")
            except Exception as e:
                st.error(f"Fout bij het aanmaken van de Opportunity: {e}")


# Opgeslagen Offertes tab



with tab2:

       
    # Knoppen voor verwijdering en vernieuwen
    col1, col2, col3 = st.columns(3)
    with col1:
        # Voeg een tabel toe met de kolommen voor "Verwerk in SAP"
        st.subheader("Verwerk in SAP")
        if "loaded_offer_df" in st.session_state:
            sap_columns = ["Artikelnummer", "Aantal", "Breedte", "Hoogte", "Spacer"]
            if all(col in st.session_state.loaded_offer_df.columns for col in sap_columns):
                sap_table = st.session_state.loaded_offer_df[sap_columns]
                st.dataframe(sap_table, use_container_width=True)
            else:
                st.warning("De geladen offerte bevat niet alle benodigde kolommen voor verwerking in SAP.")
        else:
            st.warning("Laad een offerte om de gegevens te verwerken.")
    
    with col2:
        st.write("")
        st.write("")
        st.write("")
        if "loaded_offer_df" in st.session_state:
            sap_columns = ["Artikelnummer", "Aantal", "Breedte", "Hoogte", "Spacer"]
            if all(col in st.session_state.loaded_offer_df.columns for col in sap_columns):
                sap_table = st.session_state.loaded_offer_df[sap_columns]
    
                # Knop om de inhoud van de tabel te kopiëren
                if st.button("Kopieer tabel"):
                    # Kopieer alleen de inhoud (geen headers en rijnummers)
                    content_to_copy = sap_table.to_csv(index=False, header=False, sep="\t")
                    st.write("Tabelinhoud gekopieerd naar het klembord!")
                        




# Genereer een mapping van artikelnamen naar artikelnummers
article_mapping = article_table.set_index("Description")["Material"].to_dict()


with tab3:
    # Layout met twee kolommen
    col1, col2 = st.columns(2)
    
    # Linkerkolom: Tabel met synoniemen beoordelen
    with col1:
        st.markdown("### Beoordeel output AI ✨")
    
        # Controleer of offer_df beschikbaar is in sessiestatus
        if "offer_df" in st.session_state and not st.session_state.offer_df.empty:
            # Filter regels met "Source" = "interpretatie" en "GPT"
            interpretatie_rows = st.session_state.offer_df[st.session_state.offer_df["Source"].isin(["GPT", "interpretatie"])]
    
            # Houd alleen unieke rijen op basis van combinatie van kolommen
            interpretatie_rows = interpretatie_rows.drop_duplicates(subset=["Artikelnaam", "Artikelnummer", "fuzzy_match", "original_article_number"])
        else:
            interpretatie_rows = pd.DataFrame()  # Lege DataFrame als fallback
    
        if interpretatie_rows.empty:
            st.info("Er zijn geen AI regels om te beoordelen.")
        else:
            # Maak een tabel met de correcte input en gematchte waarden
            beoordeling_tabel = interpretatie_rows.copy()
            beoordeling_tabel = beoordeling_tabel[["Artikelnaam", "Artikelnummer", "fuzzy_match", "original_article_number"]].fillna("")
            beoordeling_tabel.rename(columns={
                "Artikelnaam": "Artikelnaam",
                "Artikelnummer": "Artikelnummer",
                "original_article_number": "Jouw input",
                "fuzzy_match": "Gematcht op"
            }, inplace=True)
    
            # Configureren van de AgGrid-tabel
            gb = GridOptionsBuilder.from_dataframe(beoordeling_tabel)
            
            # Instellen van een dropdown voor de kolom "Artikelnaam"
            gb.configure_column(
                "Artikelnaam",
                editable=True,
                cellEditor="agSelectCellEditor",
                cellEditorParams={"values": list(article_mapping.keys())}
            )
            
            # Configureren van de overige kolommen
            gb.configure_column("Artikelnummer", editable=False, hide=True)
            gb.configure_column("Gematcht op", editable=False)
            gb.configure_column("Jouw input", editable=False)
            
            gb.configure_selection(selection_mode="multiple", use_checkbox=True)
            grid_options = gb.build()
    
            # Render de AgGrid-tabel
            response = AgGrid(
                beoordeling_tabel,
                gridOptions=grid_options,
                update_mode=GridUpdateMode.VALUE_CHANGED,
                fit_columns_on_grid_load=True,
                theme="material"
            )
    
            # Verwerken van wijzigingen in de tabel
            updated_rows = response["data"]  # Haal de bijgewerkte data op
            for index, row in updated_rows.iterrows():
                # Bijwerken van het artikelnummer op basis van de geselecteerde artikelnaam
                if row["Artikelnaam"] in article_mapping:
                    updated_rows.at[index, "Artikelnummer"] = article_mapping[row["Artikelnaam"]]
    
            # Knop voor accordering
            if st.button("Accordeer synoniem"):
                geselecteerde_rijen = response.get("selected_rows", pd.DataFrame())  # Haal geselecteerde rijen op als DataFrame
    
                # Controleer of de DataFrame niet leeg is
                if not geselecteerde_rijen.empty:
                    # Converteer de DataFrame naar een lijst van dictionaries
                    geselecteerde_rijen_lijst = geselecteerde_rijen.to_dict("records")
                    st.write("Geconverteerde geselecteerde rijen:", geselecteerde_rijen_lijst)  # Debug output
    
                    # Maak databaseverbinding
                    conn = create_connection()
                    cursor = conn.cursor()
    
                    try:
                        # Zorg dat de tabel SynoniemenAI bestaat
                        cursor.execute("""
                        CREATE TABLE IF NOT EXISTS SynoniemenAI (
                            Synoniem TEXT PRIMARY KEY,
                            Artikelnummer TEXT NOT NULL,
                            Artikelnaam TEXT,
                            Input TEXT,
                            Bron TEXT,
                            Datum TEXT DEFAULT CURRENT_TIMESTAMP
                        );
                        """)
    
                        # Verwerk elke rij in de lijst
                        for rij in geselecteerde_rijen_lijst:
                            input_waarde = rij.get("Jouw input", "")
                            artikelnummer = rij.get("Artikelnummer", "")
                            artikelnaam = rij.get("Artikelnaam", "")
    
                            if input_waarde and artikelnummer:
                                cursor.execute("""
                                INSERT OR IGNORE INTO SynoniemenAI (Synoniem, Artikelnummer, Artikelnaam, Input, Bron)
                                VALUES (?, ?, ?, ?, ?);
                                """, (input_waarde, artikelnummer, artikelnaam, input_waarde, "Accordeer Synoniem"))
                                st.success(f"Synoniem '{input_waarde}' -> '{artikelnummer}' is opgeslagen!")
    
                        # Commit wijzigingen naar de database
                        conn.commit()
    
                    except Exception as e:
                        st.error(f"Fout bij het opslaan: {e}")
    
                    finally:
                        # Sluit de verbinding
                        conn.close()
                else:
                    st.warning("Selecteer minimaal één rij om te accorderen of controleer de structuur.")
    


# Rechterkolom: Excel-file uploader in een expander
with col2:
    st.markdown("### Upload synoniemen 🧍‍♂⬌🧍‍♂️")
   
    with st.expander("Upload Synoniemen via Excel ✨"):
        st.markdown("Upload een Excel-bestand met de kolommen: **Artikelnummer** en **Synoniem**.")
    
        uploaded_file = st.file_uploader("Upload een Excel-bestand", type=["xlsx"])
        if uploaded_file is not None:
            try:
                # Lees de geüploade Excel-bestand
                df_synoniemen = pd.read_excel(uploaded_file)
  
                # Controleer of het bestand de juiste kolommen heeft
                if "Artikelnummer" in df_synoniemen.columns and "Synoniem" in df_synoniemen.columns:
                    if st.button("Upload🔥"):
                        # Maak een sessie aan voor SharePoint
                        session = requests.Session()
                        session.auth = HTTPBasicAuth(SP_USERNAME, SP_PASSWORD)
                        
                        headers = {
                            "Accept": "application/json;odata=verbose",
                            "Content-Type": "application/json"
                        }
                        
                        success_count = 0
                        error_count = 0

                        # Verwerk elke rij in het bestand
                        for _, row in df_synoniemen.iterrows():
                            artikelnummer = str(row["Artikelnummer"]).strip()
                            synoniem = str(row["Synoniem"]).strip()
                            gebruiker = SP_USERNAME
                            datum = pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
                            
                            # Data payload om naar SharePoint te sturen
                            data = {
                                "__metadata": {"type": "SP.Data.SynoniemenDatabaseListItem"},  # Vervang 'SynoniemenDatabaseListItem' door je juiste lijst type
                                "Artikelnummer": artikelnummer,
                                "Synoniem": synoniem,
                                "Gebruiker": gebruiker,
                                "Datum": datum
                            }
                            
                            try:
                                # URL voor het toevoegen van items aan de lijst
                                post_url = f"{SP_SITE}/_api/web/lists/getbytitle('{SP_LIST}')/items"
                                response = session.post(post_url, headers=headers, json=data)
                                
                                if response.status_code == 201:
                                    success_count += 1
                                else:
                                    error_count += 1
                            except Exception as e:
                                error_count += 1
                        
                        st.write(f"✅ Succesvol toegevoegd: {success_count}")
                        st.write(f"❌ Fouten bij toevoegen: {error_count}")
                else:
                    st.error("Het bestand moet de kolommen **'Artikelnummer'** en **'Synoniem'** bevatten.")
            except Exception as e:
                st.error(f"Fout bij het lezen van het bestand: {e}")



with col2:
    def generate_excel():
        """
        Genereer een Excel-bestand met twee tabbladen:
        1. "Synoniemen" met kolommen "Artikelnummer" en "Synoniem"
        2. "Bekende Artikelen" met de volledige artikelenlijst uit Articles.py
        """
        # Data voor tabblad 1
        synonyms_data = pd.DataFrame(columns=["Artikelnummer", "Synoniem"])
    
        # Data voor tabblad 2
        articles_data = pd.DataFrame(article_table)[["Material", "Description"]]   # Zet de geïmporteerde articles-lijst om naar een DataFrame
    
        # Schrijf naar een Excel-bestand met twee tabbladen
        output = BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            synonyms_data.to_excel(writer, sheet_name="Nieuwe synoniemen", index=False)
            articles_data.to_excel(writer, sheet_name="Artikelnummer lijst", index=False)
        
        output.seek(0)
        return output
    
    # Streamlit-interface
    st.markdown("### Download Excel voor synoniemen opvoer ⬇️")        
    
    # Maak het Excel-bestand beschikbaar voor download
    excel_file = generate_excel()
    st.download_button(
        label="Download Excel",
        data=excel_file,
        file_name="Artikelen.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    
    # # Ophalen van gegevens
    # if st.button("Haal gegevens op"):
    #     response = session.get(list_items_url, headers=headers)
    #     if response.status_code == 200:
    #         st.success("✅ Gegevens succesvol opgehaald!")
    #         data = response.json()
    #         # Toon de gegevens in een tabel
    #         if "d" in data and "results" in data["d"]:
    #             items = data["d"]["results"]
    #             for item in items:
    #                 st.write(f"ID: {item['Id']}, Synoniem: {item.get('Synoniem', 'N/A')}")
    #         else:
    #             st.warning("⚠️ Geen gegevens gevonden.")
    #     else:
    #         st.error(f"❌ Fout bij ophalen van gegevens: {response.status_code}, {response.text}")
    
# with tab5:
#     st.subheader("💬 Glasadvies Chatbot")
#     st.info("Stel je vraag over glas en krijg advies van AI op basis van beschikbare bronnen.")

    # # Functie om website content en subpagina's op te halen
    # def fetch_website_and_subpages(base_url, max_depth=0):
    #     visited_urls = set()
    #     content_list = []

    #     def crawl(url, depth):
    #         if url in visited_urls or depth > max_depth:
    #             return
    #         try:
    #             visited_urls.add(url)
    #             response = requests.get(url)
    #             if response.status_code == 200:
    #                 soup = BeautifulSoup(response.text, "html.parser")
    #                 content_list.append(soup.get_text())  # Voeg de tekst van de pagina toe
                    
    #                 # Zoek naar onderliggende links
    #                 for link in soup.find_all("a", href=True):
    #                     next_url = link["href"]
    #                     if next_url.startswith("/") or next_url.startswith(base_url):
    #                         full_url = next_url if next_url.startswith("http") else f"{base_url.rstrip('/')}/{next_url.lstrip('/')}"
    #                         crawl(full_url, depth + 1)
    #         except Exception as e:
    #             st.error(f"Fout bij ophalen van {url}: {e}")
    
    #     crawl(base_url, 0)
    #     return "\n".join(content_list)

    # # Functie om PDF-inhoud op te halen
    # def fetch_pdf_content(url):
    #     try:
    #         response = requests.get(url)
    #         pdf_file = io.BytesIO(response.content)
    #         pdf_reader = PdfReader(pdf_file)
    #         text = ""
    #         for page in pdf_reader.pages:
    #             text += page.extract_text()
    #         return text
    #     except Exception as e:
    #         st.error(f"Kon de PDF {url} niet verwerken: {e}")
    #         return ""
    
    # # Bronnen ophalen (websites + PDF)
    # sources = [
    #     fetch_website_and_subpages("https://www.onderhoudnl.nl/glasvraagbaak", max_depth=0),
    #     fetch_website_and_subpages("https://www.glasdiscount.nl/kennisbank/begrippen", max_depth=0),
    #     fetch_pdf_content("https://www.kenniscentrumglas.nl/wp-content/uploads/Infosheet-NEN-2608-1.pdf"),
    #     fetch_pdf_content("https://www.kenniscentrumglas.nl/wp-content/uploads/KCG-infosheet-Letselveiligheid-glas-NEN-3569-1.pdf"),
    # ]
    # combined_source_text = "\n".join(sources)
    
    # Initialiseer chatgeschiedenis in sessiestatus
    # if "chat_history" not in st.session_state:
    #     st.session_state["chat_history"] = [{"role": "assistant", "content": "Hoe kan ik je helpen met glasadvies?"}]
    
    # st.title("💬 Glasadvies Chatbot")
    
    # # Toon chatgeschiedenis
    # for msg in st.session_state["chat_history"]:
    #     st.chat_message(msg["role"]).write(msg["content"])
    
    # # Inputveld voor gebruikersvraag
    # user_query = st.chat_input("Stel je vraag hier:")
    
    # if user_query:
    #     st.chat_message("user").write(user_query)  # Toon de gebruikersvraag
    #     st.session_state["chat_history"].append({"role": "user", "content": user_query})
    
    #     try:
    #         # # Verstuur de vraag naar OpenAI met de opgehaalde documentatie
    #         # response = openai.chat.completions.create(
    #         #     model="gpt-4",
    #         #     messages=[
    #         #         {"role": "system", "content": "Je bent een glasadvies assistent die technisch advies geeft op basis van de gegeven documentatie. Geef kort en helder advies."},
    #         #         {"role": "user", "content": f"Documentatie:\n{combined_source_text}\n\nVraag: {user_query}"}
    #         #     ],
    #         #     max_tokens=300,
    #         #     temperature=0.7
    #         # )

    #         # # Toon het antwoord van OpenAI
    #         # ai_response = response.choices[0].message.content
    #         # st.chat_message("assistant").write(ai_response)
    #         # st.session_state["chat_history"].append({"role": "assistant", "content": ai_response})
    #         pass  # Deze logica wordt niet uitgevoerd
    #     except Exception as e:
    #         st.error(f"Er is een fout opgetreden bij het raadplegen van OpenAI: {e}")

            
